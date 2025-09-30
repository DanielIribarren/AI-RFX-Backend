"""
📄 Document Parser Service - Deserialización minimalista a texto completo
Soporta: PDF, DOCX, XLSX, TXT
Objetivo: Convertir documentos a TEXTO ÍNTEGRO sin extracciones ni adivinanzas
El LLM (Orchestrator + Analyst) hace toda la extracción inteligente
"""
import re
import logging
from typing import Dict, Any, Optional
from datetime import datetime
from pathlib import Path

logger = logging.getLogger(__name__)


class DocumentParser:
    """
    Parser MINIMALISTA que solo deserializa documentos a texto completo
    NO extrae campos, NO adivina valores, NO hace chunking
    Entrega TODO el texto al Orchestrator para extracción inteligente
    """
    
    def __init__(self):
        # Solo para limpieza mínima de caracteres de control
        self.control_chars_pattern = re.compile(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]')
    
    def parse_document(
        self, 
        content: bytes, 
        filename: str, 
        mime_type: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Deserializa documento a texto completo sin extracciones
        
        Returns:
            {
                'name': str,
                'mime': str,
                'format': str,
                'text': str,  # TEXTO ÍNTEGRO
                'char_count': int,
                'word_count': int,
                'metadata': {...}
            }
        """
        try:
            logger.info(f"📄 Deserializing document: {filename} ({mime_type})")
            
            # Detectar tipo de documento
            file_ext = Path(filename).suffix.lower()
            
            if file_ext == '.pdf' or 'pdf' in (mime_type or ''):
                return self._parse_pdf(content, filename, mime_type)
            elif file_ext in ['.docx', '.doc'] or 'wordprocessingml' in (mime_type or ''):
                return self._parse_docx(content, filename, mime_type)
            elif file_ext in ['.xlsx', '.xls'] or 'spreadsheetml' in (mime_type or ''):
                return self._parse_excel(content, filename, mime_type)
            elif file_ext == '.txt' or 'text/plain' in (mime_type or ''):
                return self._parse_text(content, filename, mime_type)
            else:
                # Fallback: intentar como texto
                logger.warning(f"⚠️ Unknown format {file_ext}, trying as text")
                return self._parse_text(content, filename, mime_type)
                
        except Exception as e:
            logger.error(f"❌ Error deserializing document {filename}: {e}", exc_info=True)
            return {
                'name': filename,
                'mime': mime_type or 'unknown',
                'format': 'error',
                'text': f'[Error deserializing document: {str(e)}]',
                'char_count': 0,
                'word_count': 0,
                'metadata': {'error': str(e), 'parsed_at': datetime.utcnow().isoformat()}
            }
    
    def _parse_pdf(self, content: bytes, filename: str, mime_type: str) -> Dict[str, Any]:
        """Deserializa PDF a texto completo con marcadores de página"""
        try:
            import PyPDF2
            from io import BytesIO
            
            pdf_file = BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            full_text = ""
            for page_num, page in enumerate(pdf_reader.pages, 1):
                page_text = page.extract_text() or ""
                full_text += f"\n\n=== Página {page_num} ===\n\n{page_text}"
            
            # Limpieza mínima
            clean_text = self._minimal_clean(full_text)
            
            logger.info(f"✅ PDF deserialized: {len(pdf_reader.pages)} pages, {len(clean_text)} chars")
            
            return {
                'name': filename,
                'mime': mime_type or 'application/pdf',
                'format': 'pdf',
                'text': clean_text,
                'char_count': len(clean_text),
                'word_count': len(clean_text.split()),
                'metadata': {
                    'pages': len(pdf_reader.pages),
                    'parsed_at': datetime.utcnow().isoformat()
                }
            }
            
        except ImportError:
            logger.warning("⚠️ PyPDF2 not installed, trying pdfplumber")
            try:
                import pdfplumber
                from io import BytesIO
                
                with pdfplumber.open(BytesIO(content)) as pdf:
                    full_text = ""
                    for page_num, page in enumerate(pdf.pages, 1):
                        page_text = page.extract_text() or ""
                        full_text += f"\n\n=== Página {page_num} ===\n\n{page_text}"
                    
                    clean_text = self._minimal_clean(full_text)
                    
                    logger.info(f"✅ PDF deserialized with pdfplumber: {len(pdf.pages)} pages")
                    return {
                        'name': filename,
                        'mime': mime_type or 'application/pdf',
                        'format': 'pdf',
                        'text': clean_text,
                        'char_count': len(clean_text),
                        'word_count': len(clean_text.split()),
                        'metadata': {
                            'pages': len(pdf.pages),
                            'parsed_at': datetime.utcnow().isoformat()
                        }
                    }
            except Exception as e:
                logger.error(f"❌ PDF deserialization failed: {e}")
                raise
    
    def _parse_docx(self, content: bytes, filename: str, mime_type: str) -> Dict[str, Any]:
        """Deserializa DOCX a texto completo (párrafos + tablas)"""
        try:
            from docx import Document
            from io import BytesIO
            
            doc = Document(BytesIO(content))
            
            full_text = ""
            
            # Extraer párrafos manteniendo viñetas y estructura
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    # Mantener viñetas como guiones
                    if para.style.name.startswith('List'):
                        full_text += f"- {text}\n"
                    else:
                        full_text += f"{text}\n"
            
            # Extraer tablas en formato TSV (tab-separated)
            if doc.tables:
                full_text += "\n\n=== TABLAS ===\n\n"
                for table_num, table in enumerate(doc.tables, 1):
                    full_text += f"--- Tabla {table_num} ---\n"
                    for row in table.rows:
                        row_cells = [cell.text.strip() for cell in row.cells]
                        full_text += '\t'.join(row_cells) + '\n'
                    full_text += '\n'
            
            # Limpieza mínima
            clean_text = self._minimal_clean(full_text)
            
            logger.info(f"✅ DOCX deserialized: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables, {len(clean_text)} chars")
            
            return {
                'name': filename,
                'mime': mime_type or 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'format': 'docx',
                'text': clean_text,
                'char_count': len(clean_text),
                'word_count': len(clean_text.split()),
                'metadata': {
                    'paragraphs': len(doc.paragraphs),
                    'tables': len(doc.tables),
                    'parsed_at': datetime.utcnow().isoformat()
                }
            }
            
        except Exception as e:
            logger.error(f"❌ DOCX deserialization failed: {e}", exc_info=True)
            raise
    
    def _parse_excel(self, content: bytes, filename: str, mime_type: str) -> Dict[str, Any]:
        """Deserializa Excel a texto completo (todas las hojas en formato TSV)"""
        try:
            import openpyxl
            from io import BytesIO
            
            wb = openpyxl.load_workbook(BytesIO(content), data_only=True)
            
            full_text = ""
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                full_text += f"\n\n=== Hoja: {sheet_name} ===\n\n"
                
                # Extraer todas las filas en formato TSV
                for row in sheet.iter_rows(values_only=True):
                    row_cells = [str(cell) if cell is not None else '' for cell in row]
                    # Solo agregar filas no vacías
                    if any(cell.strip() for cell in row_cells):
                        full_text += '\t'.join(row_cells) + '\n'
            
            # Limpieza mínima
            clean_text = self._minimal_clean(full_text)
            
            logger.info(f"✅ Excel deserialized: {len(wb.sheetnames)} sheets, {len(clean_text)} chars")
            
            return {
                'name': filename,
                'mime': mime_type or 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                'format': 'xlsx',
                'text': clean_text,
                'char_count': len(clean_text),
                'word_count': len(clean_text.split()),
                'metadata': {
                    'sheets': len(wb.sheetnames),
                    'sheet_names': wb.sheetnames,
                    'parsed_at': datetime.utcnow().isoformat()
                }
            }
            
        except Exception as e:
            logger.error(f"❌ Excel deserialization failed: {e}", exc_info=True)
            raise
    
    def _parse_text(self, content: bytes, filename: str, mime_type: str) -> Dict[str, Any]:
        """Deserializa texto plano"""
        try:
            # Intentar decodificar con diferentes encodings
            text = None
            for encoding in ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    text = content.decode(encoding)
                    logger.info(f"✅ Text decoded with {encoding}")
                    break
                except UnicodeDecodeError:
                    continue
            
            if text is None:
                text = content.decode('utf-8', errors='ignore')
                logger.warning("⚠️ Text decoded with errors ignored")
            
            # Limpieza mínima
            clean_text = self._minimal_clean(text)
            
            logger.info(f"✅ Text deserialized: {len(clean_text)} chars")
            
            return {
                'name': filename,
                'mime': mime_type or 'text/plain',
                'format': 'text',
                'text': clean_text,
                'char_count': len(clean_text),
                'word_count': len(clean_text.split()),
                'metadata': {
                    'parsed_at': datetime.utcnow().isoformat()
                }
            }
            
        except Exception as e:
            logger.error(f"❌ Text deserialization failed: {e}")
            raise
    
    def _minimal_clean(self, text: str) -> str:
        """
        Limpieza MÍNIMA: solo elimina caracteres de control y normaliza espacios
        NO recorta contenido semántico
        """
        # Eliminar caracteres de control (excepto \n y \t)
        text = self.control_chars_pattern.sub('', text)
        
        # Normalizar espacios múltiples (pero no saltos de línea)
        text = re.sub(r'[ \t]+', ' ', text)
        
        # Normalizar saltos de línea excesivos (máximo 3 consecutivos)
        text = re.sub(r'\n{4,}', '\n\n\n', text)
        
        return text.strip()



# Singleton factory
_document_parser_instance = None

def get_document_parser() -> DocumentParser:
    """Get global document parser instance"""
    global _document_parser_instance
    if _document_parser_instance is None:
        _document_parser_instance = DocumentParser()
    return _document_parser_instance
