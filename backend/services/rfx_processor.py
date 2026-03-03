"""
🎯 RFX Processor Service - Core business logic for RFX document processing
Extracts logic from rfx_webhook.py and improves it with better practices

REFACTORIZADO: Sistema modular con validaciones Pydantic, templates Jinja2,
modo debug con confidence scores y arquitectura extensible.
"""
import io
import re
import json
import PyPDF2
import time
import os
from datetime import datetime
from typing import List, Dict, Any, Optional, Union, Tuple
from enum import Enum
from openai import OpenAI
import zipfile
import mimetypes

# Feature flags for optional functionality
USE_OCR = os.getenv("RFX_USE_OCR", "true").lower() in {"1","true","yes","on"}
USE_ZIP = os.getenv("RFX_USE_ZIP", "true").lower() in {"1","true","yes","on"}

# Pydantic imports for robust validation
from pydantic import BaseModel, Field, validator, ValidationError

# Jinja2 imports for dynamic templates
from jinja2 import Template, Environment, BaseLoader

from backend.models.rfx_models import (
    RFXInput, RFXProcessed, RFXProductRequest, RFXType, RFXStatus,
    CompanyModel, RequesterModel, RFXHistoryEvent,
    # Legacy aliases for backwards compatibility
    ProductoRFX, TipoRFX, EstadoRFX
)
from backend.models.proposal_models import ProposalRequest, ProposalNotes
from backend.core.config import get_openai_config
from backend.core.database import get_database_client
# ✅ ELIMINADO: from backend.utils.validators import EmailValidator, DateValidator, TimeValidator
from backend.utils.text_utils import clean_json_string
from backend.core.feature_flags import FeatureFlags
from backend.services.function_calling_extractor import FunctionCallingRFXExtractor
from backend.services.catalog_search_service_sync import CatalogSearchServiceSync
from backend.services.ai_agents.rfx_orchestrator_agent import RFXOrchestratorAgent
from backend.services.product_resolution_service import ProductResolutionService
from backend.services.document_code_service import DocumentCodeService
from backend.utils.retry_decorator import retry_on_failure
from backend.exceptions import ExternalServiceError

import logging

logger = logging.getLogger(__name__)


# ============================================================================
# 🎯 MODELOS PYDANTIC PARA VALIDACIÓN ESTRUCTURADA
# ============================================================================

class ExtractionConfidence(BaseModel):
    """Modelo para tracking de confidence scores en extracciones"""
    field_name: str = Field(..., description="Nombre del campo extraído")
    confidence: float = Field(..., ge=0.0, le=1.0, description="Confidence score 0-1")
    source: str = Field(..., description="Fuente de la extracción (AI, manual, fallback)")
    metadata: Dict[str, Any] = Field(default_factory=dict, description="Metadata adicional")

class ProductExtraction(BaseModel):
    """Modelo Pydantic para validación robusta de productos extraídos"""
    nombre: str = Field(..., min_length=1, max_length=200, description="Nombre del producto")
    cantidad: float = Field(..., gt=0, le=10000, description="Cantidad del producto (permite decimales para kg/l)")
    unidad: str = Field(..., min_length=1, max_length=50, description="Unidad de medida")
    confidence: float = Field(default=0.8, ge=0.0, le=1.0, description="Confidence score")
    costo_unitario: Optional[float] = Field(default=0.0, ge=0, description="Costo unitario del producto")
    
    @validator('nombre')
    def validate_nombre(cls, v):
        """Validador para nombre de producto"""
        if v.lower() in ['null', 'none', 'undefined', '']:
            raise ValueError('Nombre de producto no puede ser nulo o vacío')
        return v.strip().title()
    
    @validator('unidad')
    def validate_unidad(cls, v):
        """Validador para unidad de medida"""
        valid_units = ['unidades', 'personas', 'kg', 'litros', 'porciones', 'servicios', 'pax']
        v_clean = v.lower().strip()
        if v_clean not in valid_units:
            # Auto-map common variants
            unit_mapping = {
                'unidad': 'unidades', 'pcs': 'unidades', 'units': 'unidades',
                'persona': 'personas', 'ppl': 'personas', 'people': 'personas',
                'litro': 'litros', 'l': 'litros', 'lts': 'litros',
                'porcion': 'porciones', 'servicio': 'servicios'
            }
            v_clean = unit_mapping.get(v_clean, 'unidades')
        return v_clean

class ChunkExtractionResult(BaseModel):
    """Modelo para resultado de extracción por chunk con confidence tracking"""
    email: Optional[str] = Field(None, description="Email extraído")
    email_empresa: Optional[str] = Field(None, description="Email de la empresa")
    nombre_solicitante: Optional[str] = Field(None, description="Nombre del solicitante")
    nombre_empresa: Optional[str] = Field(None, description="Nombre de la empresa") 
    telefono_solicitante: Optional[str] = Field(None, description="Teléfono del solicitante")
    telefono_empresa: Optional[str] = Field(None, description="Teléfono de la empresa")
    cargo_solicitante: Optional[str] = Field(None, description="Cargo del solicitante")
    tipo_solicitud: Optional[str] = Field(None, description="Tipo de solicitud")
    productos: List[ProductExtraction] = Field(default_factory=list, description="Lista de productos")
    fecha: Optional[str] = Field(None, description="Fecha de entrega")
    hora_entrega: Optional[str] = Field(None, description="Hora de entrega")
    lugar: Optional[str] = Field(None, description="Lugar del evento")
    texto_original_relevante: Optional[str] = Field(None, description="Texto original relevante")
    
    # Campos de confidence y debugging
    confidence_scores: List[ExtractionConfidence] = Field(default_factory=list)
    extraction_metadata: Dict[str, Any] = Field(default_factory=dict)
    chunk_index: int = Field(default=0, description="Índice del chunk procesado")

class ExtractionDebugInfo(BaseModel):
    """Información de debug para el modo debug"""
    chunk_count: int
    total_characters: int
    processing_time_seconds: float
    ai_calls_made: int
    retries_attempted: int
    confidence_summary: Dict[str, float]
    extraction_quality: str  # 'high', 'medium', 'low', 'fallback'

# ============================================================================
# 🏭 SISTEMA MODULAR DE EXTRACTORES
# ============================================================================

# ✅ ELIMINADO: ExtractionStrategy enum - El LLM siempre es inteligente y completo

class BaseExtractor:
    """Extractor base con funcionalidad común"""
    
    def __init__(self, debug_mode: bool = False):
        self.debug_mode = debug_mode
        self.confidence_threshold = 0.7
        
    def calculate_confidence(self, extracted_value: Any, raw_text: str) -> float:
        """Calcula confidence score basado en heurísticas"""
        if not extracted_value:
            return 0.0
        
        # Heurísticas básicas para confidence
        confidence = 0.5  # Base confidence
        
        # Boost si el valor extraído aparece directamente en el texto
        if isinstance(extracted_value, str) and extracted_value.lower() in raw_text.lower():
            confidence += 0.3
            
        # Boost por longitud/format validity 
        if isinstance(extracted_value, str) and len(extracted_value) > 3:
            confidence += 0.1
            
        return min(confidence, 1.0)

class ProductExtractor(BaseExtractor):
    """Extractor especializado para productos con validación robusta"""
    
    def extract_products(self, chunk_data: Dict[str, Any], raw_text: str) -> List[ProductExtraction]:
        """Extrae y valida productos de los datos del chunk"""
        productos_raw = chunk_data.get("productos", [])
        productos_validated = []
        
        if self.debug_mode:
            logger.debug(f"🔍 ProductExtractor: procesando {len(productos_raw)} productos raw")
        
        for i, producto_raw in enumerate(productos_raw):
            try:
                # Intentar crear ProductExtraction con validación Pydantic
                if isinstance(producto_raw, dict):
                    # Extract and clean product data
                    nombre = self._extract_product_name(producto_raw)
                    cantidad = self._extract_product_quantity(producto_raw)
                    unidad = self._extract_product_unit(producto_raw)
                    costo_unitario = self._extract_product_price(producto_raw)
                    
                    if nombre:  # Solo proceder si tenemos nombre válido
                        confidence = self.calculate_confidence(nombre, raw_text)
                        
                        producto = ProductExtraction(
                            nombre=nombre,
                            cantidad=cantidad,
                            unidad=unidad,
                            confidence=confidence,
                            costo_unitario=costo_unitario
                        )
                        productos_validated.append(producto)
                        
                        if self.debug_mode:
                            logger.debug(f"✅ Producto {i+1} validado: {producto.nombre} (confidence: {confidence:.2f})")
                    else:
                        if self.debug_mode:
                            logger.warning(f"⚠️ Producto {i+1} rechazado: nombre inválido")
                            
                elif isinstance(producto_raw, str) and producto_raw.strip():
                    # Handle simple string products
                    confidence = self.calculate_confidence(producto_raw, raw_text)
                    producto = ProductExtraction(
                        nombre=producto_raw.strip(),
                        cantidad=1,
                        unidad="unidades",
                        confidence=confidence,
                        costo_unitario=0.0
                    )
                    productos_validated.append(producto)
                    
            except ValidationError as e:
                if self.debug_mode:
                    logger.warning(f"⚠️ Validación fallida para producto {i+1}: {e}")
                continue
        
        return productos_validated
    
    def _extract_product_name(self, producto_dict: Dict[str, Any]) -> Optional[str]:
        """Extrae nombre del producto desde múltiples claves posibles"""
        name_keys = ["nombre", "name", "product", "producto", "item"]
        for key in name_keys:
            if key in producto_dict and producto_dict[key]:
                name = str(producto_dict[key]).strip()
                if name and name.lower() not in ["null", "none", "undefined", ""]:
                    return name
        return None
    
    def _extract_product_quantity(self, producto_dict: Dict[str, Any]) -> float:
        """Extrae cantidad del producto con fallback a 1"""
        qty_keys = ["cantidad", "quantity", "qty", "count", "numero"]
        for key in qty_keys:
            if key in producto_dict and producto_dict[key]:
                try:
                    return max(0.001, float(producto_dict[key]))
                except (ValueError, TypeError):
                    continue
        return 1.0
    
    def _extract_product_unit(self, producto_dict: Dict[str, Any]) -> str:
        """Extrae unidad del producto con fallback a 'unidades'"""
        unit_keys = ["unidad", "unit", "units", "medida"]
        for key in unit_keys:
            if key in producto_dict and producto_dict[key]:
                return str(producto_dict[key]).strip().lower()
        return "unidades"
    
    def _extract_product_price(self, producto_dict: Dict[str, Any]) -> float:
        """Extrae costo unitario del producto con fallback a 0.0"""
        price_keys = ["costo_unitario", "costo", "cost", "precio_unitario", "price", "unit_price", "precio"]
        for key in price_keys:
            if key in producto_dict and producto_dict[key] is not None:
                try:
                    price = float(producto_dict[key])
                    return max(0.0, price)  # Asegurar que no sea negativo
                except (ValueError, TypeError):
                    continue
        return 0.0

class SolicitanteExtractor(BaseExtractor):
    """Extractor especializado para información del solicitante"""
    
    def extract_solicitante_info(self, chunk_data: Dict[str, Any], raw_text: str) -> Dict[str, Any]:
        """Extrae información del solicitante con confidence tracking"""
        client_info = {}
        confidence_scores = []
        
        # Email extraction with confidence
        email = chunk_data.get("email_solicitante") or chunk_data.get("email")
        if email:
            confidence = self.calculate_confidence(email, raw_text)
            # Additional email validation boost
            if "@" in email and "." in email:
                confidence += 0.2
            client_info["email"] = email
            confidence_scores.append(ExtractionConfidence(
                field_name="email", confidence=min(confidence, 1.0), 
                source="AI", metadata={"format_valid": "@" in email}
            ))
        
        # Company email
        email_empresa = chunk_data.get("email_empresa")
        if email_empresa:
            confidence = self.calculate_confidence(email_empresa, raw_text)
            if "@" in email_empresa and "." in email_empresa:
                confidence += 0.2
            client_info["email_empresa"] = email_empresa
            confidence_scores.append(ExtractionConfidence(
                field_name="email_empresa", confidence=min(confidence, 1.0),
                source="AI", metadata={"format_valid": "@" in email_empresa}
            ))
        
        # Solicitante name extraction
        nombre_solicitante = chunk_data.get("nombre_solicitante")
        if nombre_solicitante:
            confidence = self.calculate_confidence(nombre_solicitante, raw_text)
            client_info["nombre_solicitante"] = nombre_solicitante
            confidence_scores.append(ExtractionConfidence(
                field_name="nombre_solicitante", confidence=confidence, source="AI"
            ))
        
        # Company name extraction
        nombre_empresa = chunk_data.get("nombre_empresa")
        if nombre_empresa:
            confidence = self.calculate_confidence(nombre_empresa, raw_text)
            client_info["nombre_empresa"] = nombre_empresa
            confidence_scores.append(ExtractionConfidence(
                field_name="nombre_empresa", confidence=confidence, source="AI"
            ))
        
        # Phone numbers
        for phone_field in ["telefono_solicitante", "telefono_empresa"]:
            phone = chunk_data.get(phone_field)
            if phone:
                confidence = self.calculate_confidence(phone, raw_text)
                client_info[phone_field] = phone
                confidence_scores.append(ExtractionConfidence(
                    field_name=phone_field, confidence=confidence, source="AI"
                ))
        
        # Position/cargo
        cargo = chunk_data.get("cargo_solicitante")
        if cargo:
            confidence = self.calculate_confidence(cargo, raw_text)
            client_info["cargo_solicitante"] = cargo
            confidence_scores.append(ExtractionConfidence(
                field_name="cargo_solicitante", confidence=confidence, source="AI"
            ))
        
        client_info["confidence_scores"] = confidence_scores
        return client_info

class EventExtractor(BaseExtractor):
    """Extractor especializado para información del evento"""
    
    def extract_event_info(self, chunk_data: Dict[str, Any], raw_text: str) -> Dict[str, Any]:
        """Extrae información del evento con validación"""
        event_info = {}
        confidence_scores = []
        
        # Date extraction
        fecha = chunk_data.get("fecha")
        if fecha:
            confidence = self.calculate_confidence(fecha, raw_text)
            # Boost confidence for date-like patterns
            if re.match(r'\d{4}-\d{2}-\d{2}', str(fecha)):
                confidence += 0.2
            event_info["fecha"] = fecha
            confidence_scores.append(ExtractionConfidence(
                field_name="fecha", confidence=min(confidence, 1.0), source="AI"
            ))
        
        # Time extraction  
        hora_entrega = chunk_data.get("hora_entrega")
        if hora_entrega:
            confidence = self.calculate_confidence(hora_entrega, raw_text)
            # Boost confidence for time-like patterns
            if re.match(r'\d{1,2}:\d{2}', str(hora_entrega)):
                confidence += 0.2
            event_info["hora_entrega"] = hora_entrega
            confidence_scores.append(ExtractionConfidence(
                field_name="hora_entrega", confidence=min(confidence, 1.0), source="AI"
            ))
        
        # Location extraction
        lugar = chunk_data.get("lugar")
        if lugar:
            confidence = self.calculate_confidence(lugar, raw_text)
            event_info["lugar"] = lugar
            confidence_scores.append(ExtractionConfidence(
                field_name="lugar", confidence=confidence, source="AI"
            ))
        
        # Request type
        tipo_solicitud = chunk_data.get("tipo_solicitud")
        if tipo_solicitud:
            confidence = self.calculate_confidence(tipo_solicitud, raw_text)
            event_info["tipo_solicitud"] = tipo_solicitud
            confidence_scores.append(ExtractionConfidence(
                field_name="tipo_solicitud", confidence=confidence, source="AI"
            ))
        
        event_info["confidence_scores"] = confidence_scores
        return event_info


class RFXProcessorService:
    """Service for processing RFX documents from PDF to structured data"""
    
    def __init__(self, catalog_search_service: Optional[CatalogSearchServiceSync] = None):
        self.openai_config = get_openai_config()
        # Desactivar reintentos automáticos del SDK - usamos nuestro backoff personalizado
        self.openai_client = OpenAI(
            api_key=self.openai_config.api_key,
            max_retries=0  # ← CRÍTICO: Desactivar reintentos automáticos del SDK
        )
        self.db_client = get_database_client()
        self.document_code_service = DocumentCodeService(self.db_client)
        
        self.debug_mode = False  # Simplificado - sin feature flags
        
        # 🛒 CATALOG SEARCH SERVICE SYNC (opcional - para enriquecimiento de productos)
        self.catalog_search = catalog_search_service
        if self.catalog_search:
            logger.info("🛒 Catalog Search Service (SYNC) enabled - products will be enriched with catalog prices")
        else:
            logger.info("⚠️ Catalog Search Service not provided - using AI predictions only")
        
        # 🚀 FUNCTION CALLING EXTRACTOR (siempre habilitado)
        self.function_calling_extractor = None
        self.rfx_orchestrator_agent = None
        self.product_resolution_service = None
        try:
            self.function_calling_extractor = FunctionCallingRFXExtractor(
                openai_client=self.openai_client,
                model=self.openai_config.model,
                debug_mode=self.debug_mode
            )
            self.rfx_orchestrator_agent = RFXOrchestratorAgent(
                openai_client=self.openai_client,
                model=self.openai_config.chat_model
            )
            self.product_resolution_service = ProductResolutionService(
                catalog_search=self.catalog_search,
                rfx_orchestrator_agent=self.rfx_orchestrator_agent,
            )
            logger.info("🚀 Function Calling Extractor initialized successfully")
        except Exception as e:
            logger.error(f"❌ Function Calling Extractor initialization failed: {e}")
            raise
        
        # Estadísticas de procesamiento para debugging
        self.processing_stats = {
            'total_documents_processed': 0,
            'chunks_processed': 0,
            'average_confidence': 0.0,
            'fallback_usage_count': 0,
            'catalog_matches': 0,
            'catalog_misses': 0
        }
        
        logger.info(f"🚀 RFXProcessorService inicializado - Debug: {self.debug_mode}")
    
    # ✅ ELIMINADO: _get_extraction_strategy() - No necesitamos estrategias múltiples
    
    def process_rfx_document(self, rfx_input: RFXInput, pdf_content: bytes) -> RFXProcessed:
        """
        Main processing pipeline: PDF → Text → AI Analysis → Structured Data
        """
        try:
            logger.info(f"🚀 Starting RFX processing for: {rfx_input.id}")
            
            # Step 1: Extract text from document
            extracted_text = self._extract_text_from_document(pdf_content)
            logger.info(f"📄 Text extracted: {len(extracted_text)} characters")
            
            # Store extracted text in RFXInput for later use
            rfx_input.extracted_content = extracted_text
            
            # Step 2: Process with AI
            raw_data = self._process_with_ai(extracted_text)
            logger.info(f"🤖 AI processing completed")
            
            # Step 3: Validate and clean data (enhanced with modular data)
            validated_data = self._validate_and_clean_data(raw_data, rfx_input.id)
            logger.info(f"✅ Data validated successfully")
            
            # ✅ CAMBIO #3: ELIMINADO - Ya no usamos modular_extractor
            # if self.modular_extractor.debug_mode and 'modular_debug_info' in raw_data:
            #     debug_info = raw_data['modular_debug_info']
            #     logger.info(f"📊 Modular processing stats: Time={debug_info['total_processing_time']:.3f}s")
            #     self.processing_stats['total_documents_processed'] += 1
            #     if debug_info['extraction_summary']['extraction_quality'] == 'fallback':
            #         self.processing_stats['fallback_usage_count'] += 1
            
            # Step 3.5: Intelligent RFX evaluation (if enabled)
            evaluation_metadata = self._evaluate_rfx_intelligently(validated_data, rfx_input.id)
            
            # Step 4: Create structured RFX object
            rfx_processed = self._create_rfx_processed(validated_data, rfx_input, evaluation_metadata)
            
            # Step 5: Save to database
            self._save_rfx_to_database(rfx_processed)
            
            # Step 6: RFX processing completed - Proposal generation will be handled separately by user request
            
            logger.info(f"✅ RFX processing completed successfully: {rfx_input.id}")
            return rfx_processed
            
        except Exception as e:
            logger.error(f"❌ RFX processing failed for {rfx_input.id}: {e}")
            raise
    
    def _extract_text_from_document(self, pdf_content: bytes) -> str:
        """Extract text content from PDF bytes or other file types"""
        try:
            logger.info(f"📄 Starting document text extraction, file size: {len(pdf_content)} bytes")
            
            # Try to detect file type from content
            if pdf_content.startswith(b'%PDF'):
                logger.info("📄 Detected PDF file")
                # This is a PDF file
                pdf_file = io.BytesIO(pdf_content)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                
                text_pages = []
                for i, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    text_pages.append(page_text)
                    logger.debug(f"📄 Page {i+1} extracted {len(page_text)} characters")
                
                full_text = "\n".join(text_pages)
                
                if not full_text.strip() or len(re.sub(r"\s+", "", full_text)) < 50:
                    # Attempt OCR for scanned PDFs
                    ocr_text = self._extract_text_with_ocr(pdf_content, kind="pdf")
                    if ocr_text.strip():
                        logger.info("🧠 Fallback OCR applied: PDF had no embedded text")
                        return ocr_text
                    logger.error("❌ No text extracted from PDF (even after OCR)")
                    raise ValueError("No text could be extracted from PDF")
                
                logger.info(f"✅ PDF extraction successful: {len(full_text)} total characters")
                logger.debug(f"📄 PDF text preview: {full_text[:500]}...")
                return full_text
                
            elif pdf_content.startswith(b'PK'):
                logger.info("📄 Detected DOCX file (ZIP-based)")
                # This looks like a DOCX file (ZIP format)
                try:
                    from docx import Document
                    
                    # Create a document from the bytes
                    doc_file = io.BytesIO(pdf_content)
                    doc = Document(doc_file)
                    
                    # Extract text from all paragraphs
                    paragraphs = []
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            paragraphs.append(paragraph.text.strip())
                    
                    # Extract text from tables
                    table_texts = []
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    table_texts.append(cell.text.strip())
                    
                    # Combine all text
                    all_text = "\n".join(paragraphs + table_texts)
                    
                    if not all_text.strip():
                        logger.error("❌ No text extracted from DOCX")
                        raise ValueError("No text could be extracted from DOCX")
                    
                    logger.info(f"✅ DOCX extraction successful: {len(all_text)} total characters")
                    logger.info(f"📊 Extracted {len(paragraphs)} paragraphs and {len(table_texts)} table cells")
                    logger.debug(f"📄 DOCX text preview: {all_text[:500]}...")
                    
                    return all_text
                    
                except ImportError:
                    logger.error("❌ python-docx not installed, falling back to text decode")
                    # Fallback to text decoding
                except Exception as e:
                    logger.error(f"❌ DOCX extraction failed: {e}, falling back to text decode")
            
            # Try to decode as text file
            logger.info("📄 Attempting text file decoding")
            try:
                text_content = pdf_content.decode('utf-8')
                if not text_content.strip():
                    raise ValueError("Text file is empty")
                logger.info(f"✅ UTF-8 text extraction successful: {len(text_content)} characters")
                logger.debug(f"📄 Text preview: {text_content[:500]}...")
                return text_content
            except UnicodeDecodeError:
                logger.warning("⚠️ UTF-8 decode failed, trying other encodings")
                # Try other encodings
                for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                    try:
                        text_content = pdf_content.decode(encoding)
                        if text_content.strip():
                            logger.info(f"✅ Text extraction successful with {encoding}: {len(text_content)} characters")
                            logger.debug(f"📄 Text preview: {text_content[:500]}...")
                            return text_content
                    except UnicodeDecodeError:
                        continue
                        
                logger.error("❌ All text decoding attempts failed")
                raise ValueError("Unable to decode file content")
            
        except Exception as e:
            logger.error(f"❌ Text extraction failed: {e}")
            # 🔍 Debug: Log more details about the failure
            logger.error(f"📄 File size: {len(pdf_content)} bytes")
            logger.error(f"📄 File starts with: {pdf_content[:50]}")
            logger.error(f"📄 Error type: {type(e).__name__}: {str(e)}")
            raise ValueError(f"Failed to extract text from file: {e}")
    
    def _process_with_ai(self, text: str) -> Dict[str, Any]:
        """🚀 REFACTORIZADO: Process COMPLETE text with single AI call - NO CHUNKING"""
        try:
            start_time = time.time()
            word_count = len(text.split())
            estimated_tokens = int(word_count * 1.2)  # Conservative estimate
            
            logger.info(f"🤖 Processing RFX document ({len(text)} chars, ~{estimated_tokens} tokens)")
            
            # ✅ CAMBIO #2: GPT-4o tiene 128K tokens - suficiente para documentos completos
            if estimated_tokens > 120000:
                logger.warning(f"⚠️ Document very large: {estimated_tokens} tokens - processing anyway")
                # GPT-4o puede manejar hasta 128k tokens
            
            # 🎯 ESTRATEGIA DE EXTRACCIÓN CON FALLBACKS
            extracted_data = None
            
            # 🎯 SIMPLIFICADO: Siempre usar function calling (sin feature flags)
            if self.function_calling_extractor:
                logger.info("🚀 Function Calling extraction")
                try:
                    # Obtener resultado directo de function calling (formato BD)
                    db_result = self.function_calling_extractor.extract_rfx_data(text)
                    
                    # 🔄 Convertir formato BD a formato legacy esperado
                    extracted_data = self._convert_db_result_to_legacy_format(db_result)
                    
                except Exception as e:
                    logger.error(f"❌ Function calling failed: {e}")
                    raise ExternalServiceError(
                        service_name="openai",
                        message="RFX extraction failed during function calling",
                        original_error=e,
                    )
            else:
                logger.error("❌ No function calling extractor available")
                raise ExternalServiceError(
                    service_name="openai",
                    message="Function calling extractor not available",
                )
            
            # Si no hay datos útiles, tratarlo como error de extracción para evitar falsos éxitos.
            if not extracted_data:
                logger.error("❌ All extraction methods failed")
                raise ExternalServiceError(
                    service_name="openai",
                    message="No extraction data returned by function calling",
                )
            
            # 🔍 Validate completeness automatically (for metrics only)
            completeness_result = self._validate_product_completeness(extracted_data, text)
            
            # Calculate processing metrics
            processing_time = time.time() - start_time
            self.processing_stats['total_documents_processed'] += 1
            
            # ✅ Logging simple y útil
            products_found = len(extracted_data.get('productos', []))
            logger.info(f"✅ Extraction completed in {processing_time:.2f}s - {products_found} products")
                
            # Add completeness validation results
            extracted_data['completeness_validation'] = completeness_result
            extracted_data['processing_metrics'] = {
                'processing_time': processing_time,
                'estimated_tokens': estimated_tokens,
                'extraction_method': 'function_calling'
            }
            
            return extracted_data
            
        except ExternalServiceError:
            raise
        except Exception as e:
            logger.error(f"❌ COMPLETE AI processing failed: {e}")
            raise ExternalServiceError(
                service_name="openai",
                message="AI extraction pipeline failed",
                original_error=e,
            )
            
    def _convert_db_result_to_legacy_format(self, db_result: Dict[str, Any]) -> Dict[str, Any]:
        """
        🔄 Convierte resultado de function calling en formato BD al formato legacy esperado
        
        Args:
            db_result: Resultado de function calling en formato BD
            
        Returns:
            Dict en formato legacy compatible con el resto del sistema
        """
        try:
            # Extraer datos de las diferentes secciones
            rfx_data = db_result.get('rfx_data', {})
            products_data = db_result.get('products_data', [])
            company_data = db_result.get('company_data', {})
            requester_data = db_result.get('requester_data', {})
            
            # Convertir productos del formato BD al formato legacy
            productos_legacy = []
            for product in products_data:
                raw_specs = product.get("specifications")
                specs_obj = raw_specs if isinstance(raw_specs, dict) else {}
                specs_text = ""
                if isinstance(raw_specs, str):
                    specs_text = raw_specs.strip()
                elif isinstance(specs_obj.get("details"), str):
                    specs_text = str(specs_obj.get("details")).strip()

                description = product.get("description", "") or ""
                if not description and specs_text:
                    # Preservar señal textual útil para inferencia de bundles.
                    description = specs_text

                producto_legacy = {
                    "nombre": product.get('product_name', ''),
                    "descripcion": description,
                    "categoria": product.get('category', 'otro'),
                    "cantidad": product.get('quantity', 1),
                    "unidad": product.get('unit_of_measure', 'unidades'),
                    "costo_unitario": product.get('unit_cost', 0.0),  # ✅ COSTOS UNITARIOS
                    "especificaciones": specs_obj or specs_text,
                    "es_obligatorio": product.get('is_mandatory', True),
                    "orden_prioridad": product.get('priority_order', len(productos_legacy) + 1),
                    "notas": product.get('notes', '')
                }
                productos_legacy.append(producto_legacy)
            
            # Construir resultado en formato legacy
            legacy_result = {
                # Información básica del RFX
                "title": rfx_data.get('title', ''),
                "titulo": rfx_data.get('title', ''),
                "descripcion": rfx_data.get('description', ''),
                "requirements": rfx_data.get('requirements', ''),
                
                # Fechas
                "fecha": rfx_data.get('delivery_date', ''),
                "fecha_entrega": rfx_data.get('delivery_date', ''),
                "hora_entrega": rfx_data.get('delivery_time', ''),
                
                # Ubicación
                "lugar": rfx_data.get('event_location', ''),
                "ciudad": rfx_data.get('event_city', ''),
                "pais": rfx_data.get('event_country', 'Mexico'),
                
                # Moneda y presupuesto
                "currency": rfx_data.get('currency', 'USD'),
                "presupuesto_min": rfx_data.get('budget_range_min'),
                "presupuesto_max": rfx_data.get('budget_range_max'),
                
                # Información de empresa
                "nombre_empresa": company_data.get('company_name', ''),
                "email_empresa": company_data.get('company_email', ''),
                "telefono_empresa": company_data.get('phone', ''),
                "direccion_empresa": company_data.get('address', ''),
                
                # Información de solicitante (persona)
                "nombre_solicitante": requester_data.get('name', ''),
                "email_solicitante": requester_data.get('email', ''),
                "telefono_solicitante": requester_data.get('phone', ''),
                "cargo_solicitante": requester_data.get('position', ''),
                "departamento_solicitante": requester_data.get('department', ''),
                
                # Productos (CRÍTICO - en formato legacy)
                "productos": productos_legacy,
                
                # Metadatos
                "extraction_method": "function_calling",
                "confidence_scores": rfx_data.get('metadata_json', {}).get('extraction_confidence', {}),
                "texto_original_relevante": rfx_data.get('metadata_json', {}).get('additional_metadata', {}).get('original_text_sample', ''),
                
                # Alias para compatibilidad
                "email": requester_data.get('email', ''),  # Alias para email_solicitante
            }
            
            logger.info(f"🔄 DB format converted to legacy: {len(productos_legacy)} products")
            logger.info(f"💰 Costos unitarios preserved: {sum(1 for p in productos_legacy if p.get('costo_unitario', 0) > 0)} products with costs")
            
            return legacy_result
            
        except Exception as e:
            logger.error(f"❌ Error converting DB result to legacy format: {e}")
            # Retornar formato básico en caso de error
            return {
                "productos": [],
                "email": "",
                "fecha_entrega": "",
                "hora_entrega": "",
                "lugar": "",
                "nombre_solicitante": "",
                "currency": "USD",
                "conversion_error": str(e)
            }
    
    def _robust_json_clean(self, raw_output: str) -> str:
        """🔧 Limpieza robusta de JSON que maneja markdown fences correctamente"""
        if not raw_output:
            return ""
        
        # Paso 1: Quitar markdown fences de cualquier tipo
        text = raw_output.strip()
        
        # Remover fences comunes: ```json, ```JSON, ```Json, etc.
        import re
        text = re.sub(r'^```(?:json|JSON|Json)?\s*', '', text, flags=re.MULTILINE)
        text = re.sub(r'\s*```\s*$', '', text, flags=re.MULTILINE)
        
        # Paso 2: Buscar el JSON válido más largo
        # Encuentra el primer { y el último } balanceado
        first_brace = text.find('{')
        if first_brace == -1:
            return text  # No JSON encontrado, devolver como está
        
        # Contar llaves para encontrar el cierre balanceado
        brace_count = 0
        last_brace = -1
        
        for i in range(first_brace, len(text)):
            if text[i] == '{':
                brace_count += 1
            elif text[i] == '}':
                brace_count -= 1
                if brace_count == 0:
                    last_brace = i
                    break
        
        if last_brace == -1:
            # No se encontró cierre balanceado, usar hasta el final
            json_str = text[first_brace:]
        else:
            json_str = text[first_brace:last_brace + 1]
        
        # Paso 3: Limpiar caracteres problemáticos comunes
        json_str = json_str.strip()
        
        # Remover texto antes del primer { si hay
        if json_str.startswith('JSON:') or json_str.startswith('json:'):
            json_str = json_str[5:].strip()
        
        logger.debug(f"🧹 JSON cleaning: {len(raw_output)} chars → {len(json_str)} chars")
        return json_str
    
    
    def _is_input_incomplete(self, extracted_data: Dict[str, Any], original_text: str, completeness_result: Dict[str, Any]) -> bool:
        """🚨 NUEVA FUNCIÓN: Determina si el input es incompleto y requiere más información del usuario"""
        
        logger.info(f"🔍 DEBUGGING COMPLETENESS CHECK:")
        logger.info(f"📊 Extracted data keys: {list(extracted_data.keys()) if extracted_data else 'None'}")
        
        if not extracted_data:
            logger.warning(f"❌ INCOMPLETE: No data extracted")
            return True
        
        productos = extracted_data.get('productos', [])
        logger.info(f"📦 Products to validate: {len(productos)}")
        
        # 1. NO PRODUCTS FOUND
        if not productos:
            logger.warning(f"❌ INCOMPLETE: No products found")
            return True
        
        # 2. GENERIC/VAGUE PRODUCTS - STRICTER VALIDATION
        generic_terms = ['material', 'producto', 'servicio', 'item', 'artículo', 'elemento', 'pop']
        vague_products = []
        
        for i, producto in enumerate(productos):
            nombre = producto.get('nombre', '').lower()
            descripcion = producto.get('descripcion', '').lower()
            cantidad = producto.get('cantidad', 0)
            
            logger.info(f"📦 Product {i+1}: '{nombre}' - Qty: {cantidad} - Desc: '{descripcion[:50]}...'")
            
            # Check if product name is too generic
            is_generic = any(term in nombre for term in generic_terms)
            
            # Check if missing critical information
            has_no_quantity = not cantidad or cantidad == 0
            has_no_description = not descripcion or len(descripcion.strip()) < 10
            
            # More strict: Consider vague if generic OR lacks details
            is_vague = is_generic or has_no_quantity or has_no_description
            
            if is_vague:
                vague_products.append(nombre)
                logger.warning(f"❌ VAGUE PRODUCT: '{nombre}' - Generic: {is_generic}, No quantity: {has_no_quantity}, No description: {has_no_description}")
        
        # If ANY product is vague, consider incomplete
        if vague_products:
            logger.warning(f"❌ INCOMPLETE: {len(vague_products)}/{len(productos)} products are vague: {vague_products}")
            return True
        
        # 3. MISSING ESSENTIAL INFORMATION
        essential_fields = ['nombre_solicitante', 'fecha_entrega']
        missing_fields = []
        
        for field in essential_fields:
            value = extracted_data.get(field, '')
            if not value or (isinstance(value, str) and value.strip() == ''):
                missing_fields.append(field)
        
        if missing_fields:
            logger.warning(f"❌ INCOMPLETE: Missing essential fields: {missing_fields}")
            return True
        
        # 4. COMPLETENESS RATIO TOO LOW - More strict
        completeness_ratio = completeness_result.get('completeness_ratio', 0)
        if completeness_ratio < 0.7:  # Increased from 30% to 70%
            logger.warning(f"❌ INCOMPLETE: Completeness ratio too low: {completeness_ratio:.2f}")
            return True
        
        # 5. TEXT TOO SHORT (likely incomplete description)
        if len(original_text.strip()) < 200:  # Increased from 100 to 200 characters
            logger.warning(f"❌ INCOMPLETE: Input text too short: {len(original_text)} characters")
            return True
        
        logger.info(f"✅ INPUT COMPLETE: All validation checks passed")
        return False

    def _validate_product_completeness(self, extracted_data: Dict[str, Any], original_text: str) -> Dict[str, Any]:
        """🔍 NUEVA FUNCIÓN: Valida que se extrajeron todos los productos usando heurísticas"""
        try:
            productos_found = len(extracted_data.get('productos', []))
            
            # Simple heuristics to estimate expected products
            text_lower = original_text.lower()
            
            # Count potential food/service indicators
            food_indicators = ['tequeños', 'empanadas', 'canapés', 'brochetas', 'ceviche', 'pollo', 'res', 'lasaña', 'tortas', 'cheesecake', 'frutas']
            service_indicators = ['meseros', 'montaje', 'agua', 'jugos', 'café', 'té', 'refrescos']
            
            estimated_products = 0
            for indicator in food_indicators + service_indicators:
                if indicator in text_lower:
                    estimated_products += 1
            
            completeness_ratio = productos_found / max(estimated_products, 1) if estimated_products > 0 else 1.0
            
            validation_result = {
                'products_found': productos_found,
                'estimated_products': estimated_products,
                'completeness_ratio': min(completeness_ratio, 1.0),
                'validation_status': 'complete' if completeness_ratio >= 0.8 else 'partial'
            }
            
            logger.info(f"🔍 Completeness validation: Found {productos_found}, Estimated {estimated_products}, Ratio {completeness_ratio:.2f}")
            
            return validation_result
            
        except Exception as e:
            logger.warning(f"⚠️ Completeness validation failed: {e}")
            return {'validation_status': 'unknown', 'products_found': 0}
    
    def _process_with_ai_chunked_fallback(self, text: str) -> Dict[str, Any]:
        """
        ⚠️ DEPRECATED: Ya no se usa chunking - GPT-4o tiene 128k tokens
        ✅ CAMBIO #2: Método mantenido solo por compatibilidad
        """
        logger.warning(f"⚠️ DEPRECATED: Chunked fallback called - should not be used")
        
        try:
            # Very simple fallback - try to extract at least basic info
            logger.warning(f"⚠️ Implementing simple fallback extraction")
            return self._process_with_ai_legacy(text)
            
        except Exception as e:
            logger.error(f"❌ Even fallback processing failed: {e}")
            return self._get_empty_extraction_result()
    
    def _process_with_ai_legacy(self, text: str) -> Dict[str, Any]:
        """🔧 SIMPLIFIED FALLBACK: No chunking, basic extraction only"""
        logger.warning(f"⚠️ Using SIMPLIFIED LEGACY fallback - no chunking")
        
        try:
            # Very basic extraction with JSON mode compatible prompt
            system_prompt = """Eres un extractor de datos básico. Responde ÚNICAMENTE en formato JSON válido sin markdown fences."""
            
            user_prompt = f"""Extrae información básica de este texto de catering en formato JSON:

Estructura requerida:
{{
    "title": "título corporativo relacionado a la solicitud",
    "nombre_empresa": "nombre de la empresa si se encuentra",
    "nombre_solicitante": "nombre de la persona que solicita",
    "email": "email encontrado", 
    "productos": [
        {{"nombre": "producto", "cantidad": 1, "unidad": "unidades"}}
    ],
    "fecha": "fecha si se encuentra",
    "lugar": "ubicación si se encuentra",
    "currency": "USD"
}}

TEXTO: {text[:5000]}"""

            response = self.openai_client.chat.completions.create(
                model="gpt-3.5-turbo",  # Use simpler model for fallback
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.3,
                max_tokens=1000,
                response_format={"type": "json_object"},  # 🆕 JSON MODE: Fallback legacy
                timeout=30
            )
            
            output = response.choices[0].message.content.strip()
            json_str = self._robust_json_clean(output)
            result = json.loads(json_str)
            
            # Map to expected format
            return {
                "title": result.get("title", ""),
                "email": result.get("email", ""),
                "nombre_solicitante": result.get("nombre_solicitante", ""),
                "nombre_empresa": result.get("nombre_empresa", ""), 
                "productos": result.get("productos", []),
                "fecha": result.get("fecha", ""),
                "lugar": result.get("lugar", ""),
                "currency": "USD",
                "texto_original_relevante": text[:500] if text else "",
                "processing_error": "Used simplified fallback"
            }
            
        except Exception as e:
            logger.error(f"❌ Legacy fallback failed: {e}")
            return {
                "title": "",
                "email": "",
                "nombre_solicitante": "",
                "productos": [],
                "fecha": "",
                "lugar": "",
                "currency": "USD",
                "texto_original_relevante": text[:500] if text else "",
                "processing_error": f"Legacy fallback failed: {str(e)}"
            }
    
    def _safe_get_requirements(self, validated_data: Dict[str, Any]) -> tuple:
        """Safely get requirements with fallback"""
        try:
            requirements = validated_data.get("requirements")
            confidence = validated_data.get("requirements_confidence", 0.0)
            
            # Si no hay requirements, usar valores seguros
            if not requirements or requirements in ["null", "None", ""]:
                return None, 0.0
                
            return requirements, float(confidence) if confidence else 0.0
        except Exception as e:
            logger.warning(f"⚠️ Requirements extraction failed: {e}")
            return None, 0.0
    
    def _validate_basic_requirements(self, requirements: str, confidence: float) -> Dict[str, Any]:
        """Validación mínima para MVP de requirements"""
        if not requirements:
            return {
                'validated_requirements': None,
                'adjusted_confidence': 0.0,
                'validation_issues': [],
                'needs_review': False
            }
        
        issues = []
        adjusted_confidence = confidence
        cleaned_requirements = requirements.strip()
        
        # Validación 1: Longitud apropiada
        if len(cleaned_requirements) < 5:
            issues.append("requirements_too_short")
            adjusted_confidence *= 0.3
            logger.warning(f"⚠️ Requirements muy cortos: '{cleaned_requirements}'")
        
        if len(cleaned_requirements) > 1500:
            issues.append("requirements_too_long")
            adjusted_confidence *= 0.5
            cleaned_requirements = cleaned_requirements[:1500] + "..."
            logger.warning(f"⚠️ Requirements truncados por longitud")
        
        # Validación 2: Detectar si son demasiado genéricos
        generic_phrases = [
            "necesitamos catering", "queremos evento", "solicito servicio",
            "buen servicio", "servicio de calidad", "excelente atención"
        ]
        
        requirements_lower = cleaned_requirements.lower()
        has_generic = any(phrase in requirements_lower for phrase in generic_phrases)
        
        # Detectar si hay requirements específicos
        specific_indicators = [
            "preferimos", "debe", "sin", "con experiencia", "mínimo", "máximo",
            "años", "alergia", "vegetariano", "vegano", "halal", "kosher",
            "presupuesto", "horario", "personal", "certificación"
        ]
        
        has_specific = any(indicator in requirements_lower for indicator in specific_indicators)
        
        if has_generic and not has_specific:
            adjusted_confidence *= 0.4
            issues.append("seems_too_generic")
            logger.warning(f"⚠️ Requirements parecen demasiado genéricos: {confidence:.2f} → {adjusted_confidence:.2f}")
        
        # Validación 3: Detectar si contiene información de empresa/contacto (error de extracción)
        contact_indicators = ["@", "tel:", "email:", "teléfono", "contacto:"]
        if any(indicator in requirements_lower for indicator in contact_indicators):
            adjusted_confidence *= 0.3
            issues.append("contains_contact_info")
            logger.warning(f"⚠️ Requirements contienen información de contacto")
        
        # Validación 4: Score demasiado bajo necesita revisión
        needs_review = adjusted_confidence < 0.4
        
        if needs_review:
            logger.warning(f"⚠️ Requirements necesitan revisión - Confidence: {adjusted_confidence:.2f}")
        
        return {
            'validated_requirements': cleaned_requirements,
            'adjusted_confidence': round(adjusted_confidence, 3),
            'validation_issues': issues,
            'needs_review': needs_review,
            'original_confidence': confidence
        }
    
    def _log_requirements_extraction(self, rfx_id: str, validation_result: Dict) -> None:
        """Log detallado para analizar y mejorar requirements extraction"""
        log_data = {
            'rfx_id': rfx_id,
            'timestamp': datetime.now(),
            'requirements_extracted': validation_result.get('validated_requirements'),
            'confidence_score': validation_result.get('adjusted_confidence', 0.0),
            'original_confidence': validation_result.get('original_confidence', 0.0),
            'validation_issues': validation_result.get('validation_issues', []),
            'needs_review': validation_result.get('needs_review', False),
            'requirements_length': len(validation_result.get('validated_requirements', '') or '')
        }
        
        # Log básico para monitoring
        logger.info(f"📋 Requirements extracted for {rfx_id}: "
                   f"Length={log_data['requirements_length']}, "
                   f"Confidence={log_data['confidence_score']:.3f}, "
                   f"Issues={len(log_data['validation_issues'])}")
        
        # Log detallado en debug
        if logger.isEnabledFor(logging.DEBUG):
            logger.debug(f"📊 Requirements log data: {log_data}")
        
        # TODO: En el futuro, guardar en tabla de logs para análisis
        # self.requirements_logs.insert(log_data)
    
    def _validate_and_normalize_currency(self, currency: str) -> str | None:
        """🆕 Valida y normaliza códigos de moneda ISO 4217"""
        if not currency or currency.lower() in ["null", "none", "", "undefined"]:
            logger.info(f"📝 No currency provided - left as null for user completion")
            return None
        
        # Lista de monedas válidas más comunes en el contexto de negocio
        valid_currencies = {
            # Monedas principales
            "USD", "EUR", "GBP", "JPY", "CHF", "CAD", "AUD", "NZD",
            # Monedas latinoamericanas
            "MXN", "BRL", "ARS", "COP", "PEN", "CLP", "UYU", "BOB", "VES",
            # Monedas asiáticas
            "CNY", "HKD", "SGD", "KRW", "INR", "THB", "MYR", "IDR",
            # Otras monedas relevantes
            "NOK", "SEK", "DKK", "PLN", "CZK", "HUF", "RUB", "ZAR"
        }
        
        # Normalizar entrada
        currency_clean = currency.strip().upper()
        
        # Validación directa
        if currency_clean in valid_currencies:
            logger.debug(f"💰 Currency validated: {currency_clean}")
            return currency_clean
        
        # Mapeo de aliases comunes
        currency_aliases = {
            # Símbolos a códigos
            "$": "USD",  # Predeterminado para $
            "€": "EUR",
            "£": "GBP", 
            "¥": "JPY",
            "₽": "RUB",
            "₹": "INR",
            
            # Variaciones textuales
            "DOLLAR": "USD",
            "DOLARES": "USD", 
            "DÓLARES": "USD",
            "DOLLARS": "USD",
            "EURO": "EUR",
            "EUROS": "EUR",
            "POUND": "GBP",
            "POUNDS": "GBP",
            "LIBRA": "GBP",
            "LIBRAS": "GBP",
            "YEN": "JPY",
            "PESO": "MXN",  # Predeterminado para peso sin contexto
            "PESOS": "MXN",
            "REAL": "BRL",
            "REALES": "BRL",
            "R$": "BRL",
            
            # Códigos con sufijos
            "USD$": "USD",
            "CAD$": "CAD",
            "AUD$": "AUD",
            "MXN$": "MXN",
            
            # Variaciones regionales
            "DÓLAR": "USD",
            "DÓLARES AMERICANOS": "USD",
            "DÓLARES ESTADOUNIDENSES": "USD",
            "PESOS MEXICANOS": "MXN",
            "PESOS COLOMBIANOS": "COP",
            "SOLES": "PEN",
            "BOLIVARES": "VES"
        }
        
        # Buscar en aliases
        if currency_clean in currency_aliases:
            mapped_currency = currency_aliases[currency_clean]
            logger.debug(f"💰 Currency mapped: {currency} → {mapped_currency}")
            return mapped_currency
        
        # Intentar extraer código de una cadena más larga
        for valid_code in valid_currencies:
            if valid_code in currency_clean:
                logger.debug(f"💰 Currency extracted from text: {currency} → {valid_code}")
                return valid_code
        
        # Si no se puede validar, preservar para revisión manual
        logger.warning(f"⚠️ Unrecognized currency '{currency}' - preserved for manual review")
        return currency
    
    def _validate_and_clean_data(self, raw_data: Dict[str, Any], rfx_id: str) -> Dict[str, Any]:
        """
        ✅ SIMPLIFICADO: El LLM ya validó, solo verificar estructura básica y mapear campos
        """
        logger.info(f"🔍 Validating structure for RFX: {rfx_id}")
        
        # Handle None raw_data
        if raw_data is None:
            logger.warning(f"⚠️ raw_data is None for RFX: {rfx_id}")
            raw_data = {}
        
        # Log para debugging - ver qué campos vienen del LLM
        logger.debug(f"📦 Raw data keys: {list(raw_data.keys())}")
        
        # Mapear campos - El LLM retorna con nombres específicos según el prompt
        requester_name = (
            raw_data.get('nombre_solicitante')
            or raw_data.get('requester_name')
            or ''
        )
        company_name = (
            raw_data.get('nombre_empresa')
            or raw_data.get('empresa')
            or raw_data.get('company_name')
            or ''
        )
        # Evitar mezclar empresa y solicitante cuando llegan idénticos por extracción ambigua.
        if requester_name and company_name and requester_name.strip().lower() == company_name.strip().lower():
            requester_name = ''

        title_candidate = (
            raw_data.get('title')
            or raw_data.get('titulo')
            or raw_data.get('request_title')
            or ''
        )
        description_candidate = (
            raw_data.get('description')
            or raw_data.get('descripcion')
            or raw_data.get('requirements')
            or ''
        )

        validated = {
            # Título y contexto (debe venir del modelo AI)
            'title': self._sanitize_text_field(title_candidate, max_length=200),
            'description': self._sanitize_text_field(description_candidate, max_length=2000),
            'texto_original_relevante': self._sanitize_text_field(raw_data.get('texto_original_relevante', ''), max_length=500),

            # Campos del solicitante (persona)
            'email': raw_data.get('email_solicitante') or raw_data.get('email', ''),
            'nombre_solicitante': requester_name,
            'telefono_solicitante': raw_data.get('telefono_solicitante') or raw_data.get('telefono', ''),
            'cargo_solicitante': raw_data.get('cargo_solicitante') or raw_data.get('cargo', ''),
            
            # Campos de empresa (organización)
            'nombre_empresa': company_name,
            'email_empresa': raw_data.get('email_empresa') or '',
            'telefono_empresa': raw_data.get('telefono_empresa') or '',
            
            # Campos del evento
            'fecha': raw_data.get('fecha') or raw_data.get('fecha_entrega', ''),
            'hora_entrega': raw_data.get('hora_entrega') or raw_data.get('hora', ''),
            'lugar': raw_data.get('lugar', ''),
            
            # Productos y otros
            'currency': raw_data.get('currency', 'USD'),
            'productos': raw_data.get('productos', []),
            'requirements': raw_data.get('requirements', ''),
            'requirements_confidence': raw_data.get('requirements_confidence', 0.0)
        }

        # Normalización defensiva de productos: evitar ítems vacíos que rompan persistencia.
        cleaned_products: List[Dict[str, Any]] = []
        for product in validated.get("productos") or []:
            if not isinstance(product, dict):
                continue
            name = str(product.get("nombre") or product.get("product_name") or "").strip()
            if not name:
                continue
            normalized = dict(product)
            normalized["nombre"] = name
            if normalized.get("cantidad") is None:
                normalized["cantidad"] = normalized.get("quantity", 1)
            if normalized.get("unidad") is None:
                normalized["unidad"] = normalized.get("unit", "unidades")
            cleaned_products.append(normalized)
        validated["productos"] = cleaned_products
        
        # Log de campos importantes para debugging
        logger.info(f"✅ Validated: {len(validated.get('productos', []))} products, "
                   f"email={validated.get('email')}, empresa={validated.get('nombre_empresa')}")
        
        return validated

    def _sanitize_text_field(self, value: Any, max_length: int = 200) -> str:
        """Normalize text fields defensively for storage/display."""
        if value is None:
            return ""

        text = re.sub(r"\s+", " ", str(value)).strip()
        if not text:
            return ""
        return text[:max_length]

    def _is_generic_title(self, title: str) -> bool:
        """Detect generic title patterns that should not be used as final RFX title."""
        normalized = re.sub(r"\s+", " ", title.strip().lower())
        if not normalized:
            return True

        generic_exact = {
            "rfx",
            "rfq",
            "rfp",
            "rfi",
            "rfx request",
            "solicitud",
            "solicitud de presupuesto",
            "request for quote",
            "request for proposal",
            "request for information",
            "evento corporativo",
        }
        if normalized in generic_exact:
            return True

        return normalized.startswith("rfx request")

    def _build_contextual_title_fallback(self, validated_data: Dict[str, Any], rfx_input: RFXInput) -> str:
        """
        Build a contextual title only when AI title is missing/unusable.
        Avoid static hardcoded templates like 'RFX Request - ...'.
        """
        for candidate in (
            validated_data.get("description"),
            validated_data.get("requirements"),
            validated_data.get("texto_original_relevante"),
        ):
            cleaned = self._sanitize_text_field(candidate, max_length=200)
            if cleaned and len(cleaned) >= 12:
                return cleaned

        context_parts = [
            self._sanitize_text_field(validated_data.get("nombre_empresa"), max_length=80),
            self._sanitize_text_field(validated_data.get("nombre_solicitante"), max_length=80),
            self._sanitize_text_field(validated_data.get("lugar"), max_length=80),
        ]
        context_parts = [part for part in context_parts if part]
        if context_parts:
            return " | ".join(context_parts)[:200]

        rfx_type = rfx_input.rfx_type.value if hasattr(rfx_input.rfx_type, "value") else str(rfx_input.rfx_type)
        return f"Solicitud {rfx_type}".strip()[:200]

    def _resolve_rfx_title(self, validated_data: Dict[str, Any], rfx_input: RFXInput) -> str:
        """Prefer AI-generated title, with contextual fallback when missing/generic."""
        ai_title = self._sanitize_text_field(validated_data.get("title"), max_length=200)
        if ai_title and not self._is_generic_title(ai_title):
            return ai_title

        if ai_title:
            logger.warning(f"⚠️ Generic AI title detected, using contextual fallback: '{ai_title}'")
        else:
            logger.warning("⚠️ AI title missing, using contextual fallback")

        fallback_title = self._build_contextual_title_fallback(validated_data, rfx_input)
        return self._sanitize_text_field(fallback_title, max_length=200)
    
    def _create_rfx_processed(self, validated_data: Dict[str, Any], rfx_input: RFXInput, evaluation_metadata: Optional[Dict[str, Any]] = None) -> RFXProcessed:
        """Create RFXProcessed object from validated data and evaluation results"""
        try:
            # 🔍 DEBUG: Log object creation
            logger.info(f"🔨 Creating RFXProcessed object for: {rfx_input.id}")
            resolved_title = self._resolve_rfx_title(validated_data, rfx_input)
            
            # Convert productos to ProductoRFX objects (map Spanish to English fields)
            productos = [
                ProductoRFX(
                    product_name=p["nombre"],
                    quantity=p["cantidad"],
                    unit=p["unidad"],
                    costo_unitario=p.get("costo_unitario", 0.0),  # ✅ Costo del catálogo
                    precio_unitario=p.get("precio_unitario", 0.0)  # ✅ FIX: Incluir precio del catálogo
                )
                for p in validated_data["productos"]
            ]
            
            # 🔍 DEBUG: Log productos creados con precios
            logger.info(f"🔨 Created {len(productos)} ProductoRFX objects")
            if productos:
                first = productos[0]
                logger.info(f"   📦 First product: {first.product_name}")
                logger.info(f"   💰 costo_unitario: {first.costo_unitario}")
                logger.info(f"   💰 precio_unitario: {first.precio_unitario}")
            
            # Prepare enhanced metadata including empresa data
            metadata = {
                "original_rfx_id": rfx_input.id,  # Preserve original string ID for frontend
                "texto_original_length": len(rfx_input.extracted_content or ""),
                "productos_count": len(productos),
                "processing_version": "2.1",  # Upgraded version with intelligent evaluation
                "validation_status": validated_data.get("validation_metadata", {}),
                "texto_original_relevante": validated_data.get("texto_original_relevante", ""),
                "ai_extraction_quality": "high" if validated_data.get("validation_metadata", {}).get("has_original_data", False) else "fallback",
                # ✨ AÑADIR: Datos de empresa en metadatos para acceso del frontend
                "nombre_empresa": validated_data.get("nombre_empresa", ""),
                "email_empresa": validated_data.get("email_empresa", ""),
                "telefono_empresa": validated_data.get("telefono_empresa", ""),
                "telefono_solicitante": validated_data.get("telefono_solicitante", ""),
                "cargo_solicitante": validated_data.get("cargo_solicitante", ""),
                # ✨ MONEDA: Currency extraída por AI
                "validated_currency": validated_data.get("currency", "USD"),
                # 🔒 Preservar estructura rica (bundle/specifications) para persistencia final
                "validated_products_full": validated_data.get("productos", []),
            }
            
            # Integrate intelligent evaluation metadata if available
            if evaluation_metadata:
                metadata["intelligent_evaluation"] = evaluation_metadata
                
                # Log evaluation integration
                if evaluation_metadata.get('evaluation_enabled'):
                    if 'evaluation_error' in evaluation_metadata:
                        logger.warning(f"⚠️ Evaluation error included in metadata for {rfx_input.id}")
                    else:
                        score = evaluation_metadata.get('execution_summary', {}).get('consolidated_score')
                        quality = evaluation_metadata.get('execution_summary', {}).get('overall_quality')
                        domain = evaluation_metadata.get('domain_detection', {}).get('primary_domain')
                        logger.info(f"📊 Evaluation metadata integrated - Score: {score}, Quality: {quality}, Domain: {domain}")
                else:
                    logger.debug(f"🔧 Evaluation disabled - metadata reflects feature flag status")
            else:
                logger.debug(f"ℹ️ No evaluation metadata provided for {rfx_input.id}")
            
            # 🔍 DEBUG: Log metadata
            logger.debug(f"📊 Metadata prepared: {metadata}")
            
            # Generate UUID for database, but keep original ID in metadata
            from uuid import uuid4
            rfx_uuid = uuid4()
            
            rfx_processed = RFXProcessed(
                id=rfx_uuid,
                rfx_type=rfx_input.rfx_type,
                title=resolved_title,
                description=validated_data.get("description") or None,
                location=validated_data["lugar"] or None,
                delivery_date=validated_data["fecha"] or None,
                delivery_time=validated_data["hora_entrega"] or None,  # ✅ AI-FIRST: None si vacío
                status=RFXStatus.IN_PROGRESS,
                original_pdf_text=rfx_input.extracted_content,
                requested_products=[p.dict() for p in productos] if productos else [],
                metadata_json=metadata,
                
                # Legacy/extracted fields for compatibility
                email=validated_data["email"],
                requester_name=validated_data["nombre_solicitante"],
                company_name=validated_data.get("nombre_empresa", ""),
                products=productos,  # productos is already a list of RFXProductRequest objects
                
                # 🆕 MVP: Requirements específicos del cliente (SAFE)
                requirements=self._safe_get_requirements(validated_data)[0],
                requirements_confidence=self._safe_get_requirements(validated_data)[1]
            )
            
            # 🔍 DEBUG: Log successful creation
            logger.info(f"✅ RFXProcessed object created successfully")
            empresa_info = metadata or {}
            empresa_nombre = empresa_info.get("nombre_empresa", "No especificada")
            original_id = empresa_info.get("original_rfx_id", str(rfx_processed.id))
            logger.info(f"📦 RFX Object: Original ID={original_id}, UUID={rfx_processed.id}, Solicitante={rfx_processed.requester_name}, Empresa={empresa_nombre}, Productos={len(rfx_processed.products)}")
            
            # Log empresa details if available
            if empresa_info.get("nombre_empresa"):
                logger.info(f"🏢 Empresa: {empresa_info.get('nombre_empresa')} | Email: {empresa_info.get('email_empresa', 'N/A')} | Tel: {empresa_info.get('telefono_empresa', 'N/A')}")
            
            # Log solicitante details if available  
            if empresa_info.get("telefono_solicitante") or empresa_info.get("cargo_solicitante"):
                logger.info(f"👤 Solicitante: {rfx_processed.requester_name} | Tel: {empresa_info.get('telefono_solicitante', 'N/A')} | Cargo: {empresa_info.get('cargo_solicitante', 'N/A')}")
            
            # 🆕 MVP: Log requirements info if available
            if rfx_processed.requirements:
                req_preview = rfx_processed.requirements[:100] + "..." if len(rfx_processed.requirements) > 100 else rfx_processed.requirements
                logger.info(f"📋 Requirements: '{req_preview}' | Confidence: {rfx_processed.requirements_confidence:.3f}")
            else:
                logger.debug(f"📝 No requirements extracted for RFX")
                
            return rfx_processed
            
        except Exception as e:
            logger.error(f"❌ Failed to create RFXProcessed object: {e}")
            raise
    
    def _save_rfx_to_database(self, rfx_processed: RFXProcessed, user_id: str = None, organization_id: str = None) -> None:
        """Save processed RFX to database V2.0 with normalized structure"""
        try:
            if user_id:
                logger.info(f"💾 Saving RFX with user_id: {user_id}")
            else:
                logger.warning(f"⚠️ Saving RFX without user_id - will be NULL in database")
            
            if organization_id:
                logger.info(f"💾 Saving RFX with organization_id: {organization_id}")
            else:
                logger.warning(f"⚠️ Saving RFX without organization_id - will be NULL in database")
            # Extract company and requester information from metadata
            metadata = rfx_processed.metadata_json or {}
            
            # 1. Create or find company
            # Clean email to meet database constraints (no spaces allowed)
            email_empresa = metadata.get("email_empresa", "")
            if email_empresa:
                email_empresa = email_empresa.replace(" ", "").strip()
                if not email_empresa or "@" not in email_empresa:
                    email_empresa = None
            else:
                email_empresa = None
                
            company_data = {
                "name": metadata.get("nombre_empresa", "Unknown Company"),
                "email": email_empresa,
                "phone": metadata.get("telefono_empresa")
            }
            company_record = self.db_client.insert_company(company_data)
            
            # 2. Create or find requester
            # Clean requester email to meet database constraints
            requester_email = rfx_processed.email or ""
            if requester_email:
                requester_email = requester_email.replace(" ", "").strip()
                if not requester_email or "@" not in requester_email:
                    requester_email = None
            else:
                requester_email = None
                
            requester_data = {
                "company_id": company_record.get("id"),
                "name": rfx_processed.requester_name or "Unknown Requester",
                "email": requester_email,
                "phone": metadata.get("telefono_solicitante"),
                "position": metadata.get("cargo_solicitante")
            }
            requester_record = self.db_client.insert_requester(requester_data)
            
            # 3. Prepare RFX data for database (V2.0 schema)
            rfx_data = {
                "id": str(rfx_processed.id),
                "company_id": company_record.get("id"),
                "requester_id": requester_record.get("id"),
                "rfx_type": rfx_processed.rfx_type.value if hasattr(rfx_processed.rfx_type, 'value') else str(rfx_processed.rfx_type),
                "title": rfx_processed.title or f"{rfx_processed.company_name}",
                "location": rfx_processed.location,
                "delivery_date": rfx_processed.delivery_date.isoformat() if rfx_processed.delivery_date else None,
                "delivery_time": rfx_processed.delivery_time.isoformat() if rfx_processed.delivery_time else None,
                "status": rfx_processed.status.value if hasattr(rfx_processed.status, 'value') else str(rfx_processed.status),
                "original_pdf_text": rfx_processed.original_pdf_text,
                "requested_products": rfx_processed.requested_products or [],
                "metadata_json": rfx_processed.metadata_json,
                "currency": rfx_processed.metadata_json.get('validated_currency', 'USD') if rfx_processed.metadata_json else 'USD',
                
                # 🆕 MVP: Requirements específicos del cliente
                "requirements": rfx_processed.requirements,
                "requirements_confidence": rfx_processed.requirements_confidence
            }

            # Corporate-friendly deterministic code for fast lookup and classification.
            metadata_json = rfx_data.get("metadata_json") if isinstance(rfx_data.get("metadata_json"), dict) else {}
            existing_rfx_code = metadata_json.get("rfx_code")
            if existing_rfx_code:
                rfx_code = str(existing_rfx_code)
            else:
                try:
                    rfx_code = self.document_code_service.generate_rfx_code(rfx_data.get("rfx_type"))
                except Exception:
                    # Fallback only if sequence RPC is unavailable.
                    current_year = datetime.utcnow().year
                    rfx_code = f"RFX-OTH-{current_year}-{str(rfx_processed.id).split('-')[0].upper()}"
                    logger.warning(f"⚠️ Using fallback rfx_code generation: {rfx_code}")

            # Backward-compatible: keep canonical code in metadata_json.
            # Some deployed DBs may not yet have rfx_v2.rfx_code.
            metadata_json["rfx_code"] = rfx_code
            rfx_data["metadata_json"] = metadata_json
            logger.info(f"🏷️ Assigned RFX code: {rfx_code}")
            
            # 🆕 CRITICAL: Add user_id if provided
            if user_id:
                rfx_data["user_id"] = user_id
                logger.info(f"✅ Added user_id to rfx_data: {user_id}")
            else:
                logger.warning(f"⚠️ No user_id provided - rfx_data will not have user_id field")
            
            # 🆕 CRITICAL: Add organization_id if provided
            if organization_id:
                rfx_data["organization_id"] = organization_id
                logger.info(f"✅ Added organization_id to rfx_data: {organization_id}")
            else:
                logger.warning(f"⚠️ No organization_id provided - rfx_data will not have organization_id field")
            
            # 4. Insert RFX data
            rfx_record = self.db_client.insert_rfx(rfx_data)
            
            # 5. Insert structured products if available
            if rfx_processed.products:
                structured_products = []
                full_products = []
                if isinstance(rfx_processed.metadata_json, dict):
                    full_products = list(rfx_processed.metadata_json.get("validated_products_full") or [])

                for idx, product in enumerate(rfx_processed.products):
                    # ✅ FIX: Productos enriquecidos son DICT, no objetos - usar .get() en vez de getattr()
                    # Intentar múltiples nombres de campo para compatibilidad
                    if isinstance(product, dict):
                        # Producto es diccionario (enriquecido con catálogo)
                        # 🔍 DEBUG: Log del producto completo para ver qué campos tiene
                        logger.info(f"   📦 Product dict keys: {list(product.keys())}")
                        logger.info(f"   📦 Full product data: {product}")
                        
                        # Intentar obtener costo con fallback
                        costo = product.get('costo_unitario')
                        if costo is None:
                            costo = product.get('unit_cost')
                        
                        # Intentar obtener precio con fallback
                        precio = product.get('precio_unitario')
                        if precio is None:
                            precio = product.get('unit_price')
                        
                        # 🔍 DEBUG: Log de valores extraídos
                        logger.info(f"   💰 Extracted - costo_unitario: {product.get('costo_unitario')}, unit_cost: {product.get('unit_cost')}")
                        logger.info(f"   💰 Extracted - precio_unitario: {product.get('precio_unitario')}, unit_price: {product.get('unit_price')}")
                        logger.info(f"   💰 Final values - costo: {costo}, precio: {precio}")
                        logger.info(f"   📊 Catalog match: {product.get('catalog_match')}, pricing_source: {product.get('pricing_source')}")
                        
                        product_name = product.get('product_name') or product.get('nombre')
                        quantity = product.get('quantity') or product.get('cantidad')
                        unit = product.get('unit') or product.get('unidad')
                    else:
                        # Producto es objeto (legacy)
                        costo = getattr(product, 'costo_unitario', None)
                        if costo is None:
                            costo = getattr(product, 'unit_cost', None)
                        precio = getattr(product, 'precio_unitario', None)
                        if precio is None:
                            precio = getattr(product, 'unit_price', None)
                        product_name = product.product_name if hasattr(product, 'product_name') else product.nombre
                        quantity = product.quantity if hasattr(product, 'quantity') else product.cantidad
                        unit = product.unit if hasattr(product, 'unit') else product.unidad

                    full_product = full_products[idx] if idx < len(full_products) and isinstance(full_products[idx], dict) else {}
                    quantity_final = quantity if quantity is not None else full_product.get("cantidad") or full_product.get("quantity") or 1
                    unit_final = unit or full_product.get("unidad") or full_product.get("unit") or "unidades"
                    price_final = precio if precio is not None else (costo if costo is not None else None)
                    specs_final = (
                        (product.get("specifications") if isinstance(product, dict) else None)
                        or full_product.get("specifications")
                        or full_product.get("especificaciones")
                        or {}
                    )
                    if not isinstance(specs_final, dict):
                        specs_final = {}
                    
                    product_data = {
                        "product_name": product_name,
                        "quantity": quantity_final,
                        "unit": unit_final,
                        "estimated_unit_price": price_final,  # Fallback: si solo hay costo explícito, usarlo como precio comercial base.
                        "unit_cost": costo if costo is not None else None,  # ✅ FIX: Permitir 0 como valor válido
                        "description": (
                            product.get("descripcion")
                            if isinstance(product, dict)
                            else full_product.get("descripcion")
                        ),
                        "specifications": specs_final,
                        "notes": (
                            product.get("notas")
                            if isinstance(product, dict) and product.get("notas")
                            else full_product.get("notas") or "Extracted from RFX processing"
                        ),
                    }
                    structured_products.append(product_data)
                    
                    # 🔍 DEBUG: Log each product's pricing
                    logger.info(f"   💰 Saving product: {product_data['product_name']} - cost: ${costo if costo is not None else 'None'}, price: ${precio if precio is not None else 'None'}")
                
                self.db_client.insert_rfx_products(rfx_record["id"], structured_products)
                logger.info(f"✅ {len(structured_products)} structured products saved")
                
                # 🔍 DEBUG: Log summary of pricing
                total_with_price = sum(1 for p in structured_products if p.get('estimated_unit_price') and p['estimated_unit_price'] > 0)
                logger.info(f"   💰 Products with estimated_unit_price > 0: {total_with_price}/{len(structured_products)}")
            
            # 6. Create history event
            history_event = {
                "rfx_id": str(rfx_processed.id),
                "event_type": "rfx_processed",
                "description": f"RFX processed successfully with {len(rfx_processed.products or [])} products",
                "new_values": {
                    "company_name": company_data["name"],
                    "requester_name": requester_data["name"],
                    "processing_version": metadata.get("processing_version", "2.0"),
                    "original_rfx_id": metadata.get("original_rfx_id"),
                    "product_count": len(rfx_processed.products or [])
                },
                "performed_by": "system_ai"
            }
            self.db_client.insert_rfx_history(history_event)
            
            logger.info(f"✅ RFX saved to database V2.0: {rfx_processed.id}")
            
        except Exception as e:
            logger.error(f"❌ Failed to save RFX to database: {e}")
            raise
    
    # REMOVED: _generate_proposal_automatically
    # La generación de propuestas ahora se maneja por separado cuando el usuario lo solicite
    # mediante el endpoint /api/proposals/generate después de revisar datos y establecer costos
    
    def _map_rfx_data_for_proposal(self, rfx_data_raw: Dict[str, Any]) -> Dict[str, Any]:
        """
        🔧 Mapea estructura BD V2.0 (inglés, normalizada) → ProposalGenerationService (español, plana)
        
        Args:
            rfx_data_raw: Datos del RFX desde BD V2.0 con companies, requesters, etc.
            
        Returns:
            Dict con estructura esperada por ProposalGenerationService
        """
        try:
            # Extraer datos de BD V2.0
            company_data = rfx_data_raw.get("companies", {}) or {}
            requester_data = rfx_data_raw.get("requesters", {}) or {}
            requested_products = rfx_data_raw.get("requested_products", [])
            metadata = rfx_data_raw.get("metadata_json", {}) or {}
            
            # Mapear productos con fallback a metadata si es necesario
            productos_mapped = []
            if requested_products:
                # Usar productos estructurados de BD
                for producto in requested_products:
                    productos_mapped.append({
                        "nombre": producto.get("product_name", producto.get("nombre", "Producto")),
                        "cantidad": producto.get("quantity", producto.get("cantidad", 1)),
                        "unidad": producto.get("unit", producto.get("unidad", "unidades"))
                    })
            else:
                # Fallback a metadata si no hay productos estructurados
                metadata_productos = metadata.get("productos", [])
                for producto in metadata_productos:
                    if isinstance(producto, dict):
                        productos_mapped.append({
                            "nombre": producto.get("nombre", "Producto"),
                            "cantidad": producto.get("cantidad", 1),
                            "unidad": producto.get("unidad", "unidades")
                        })
            
            # Si aún no hay productos, crear uno por defecto
            if not productos_mapped:
                productos_mapped = [{
                    "nombre": "Servicio de Catering",
                    "cantidad": 1,
                    "unidad": "servicio"
                }]
            
            # Estructura mapeada esperada por ProposalGenerationService
            mapped_data = {
                # ✅ Mapear información del cliente (combinando empresa + solicitante)
                "clientes": {
                    "nombre": requester_data.get("name", metadata.get("nombre_solicitante", "Cliente")),
                    "email": requester_data.get("email", metadata.get("email", "")),
                    "empresa": company_data.get("name", metadata.get("nombre_empresa", "")),
                    "telefono": requester_data.get("phone", metadata.get("telefono_solicitante", "")),
                    "cargo": requester_data.get("position", metadata.get("cargo_solicitante", "")),
                    # Información adicional de empresa
                    "email_empresa": company_data.get("email", metadata.get("email_empresa", "")),
                    "telefono_empresa": company_data.get("phone", metadata.get("telefono_empresa", ""))
                },
                
                # ✅ Mapear productos
                "productos": productos_mapped,
                
                # ✅ Mapear información del evento
                "lugar": rfx_data_raw.get("location", metadata.get("lugar", "Por definir")),
                "fecha_entrega": rfx_data_raw.get("delivery_date", metadata.get("fecha", "")),
                "hora_entrega": rfx_data_raw.get("delivery_time", metadata.get("hora_entrega", "")),
                "tipo": rfx_data_raw.get("rfx_type", metadata.get("tipo_solicitud", "catering")),
                
                # ✅ Información adicional
                "id": str(rfx_data_raw.get("id", "")),
                "title": rfx_data_raw.get("title", ""),
                "texto_original_relevante": metadata.get("texto_original_relevante", "")
            }
            
            logger.debug(f"🔧 Datos mapeados para propuesta: Cliente={mapped_data['clientes']['nombre']}, "
                        f"Empresa={mapped_data['clientes']['empresa']}, Productos={len(mapped_data['productos'])}")
            
            return mapped_data
            
        except Exception as e:
            logger.error(f"❌ Error mapeando datos para propuesta: {e}")
            # Fallback con estructura mínima
            return {
                "clientes": {"nombre": "Cliente", "email": ""},
                "productos": [{"nombre": "Servicio de Catering", "cantidad": 1, "unidad": "servicio"}],
                "lugar": "Por definir",
                "fecha_entrega": "",
                "hora_entrega": "",
                "tipo": "catering"
            }
    
    # REMOVED: _log_proposal_generation_event and _log_proposal_generation_error
    # Estas funciones ya no se necesitan porque la generación de propuestas es manual

    @retry_on_failure(max_retries=3, initial_delay=1.0, backoff_factor=2.0)
    def _call_openai_with_retry(self, **kwargs) -> Any:
        """
        Call OpenAI API with automatic retry logic.
        Uses unified retry decorator with exponential backoff.
        """
        try:
            return self.openai_client.chat.completions.create(**kwargs)
        except Exception as e:
            logger.error(f"OpenAI API call failed: {e}")
            raise ExternalServiceError("OpenAI", str(e), original_error=e)

    def _evaluate_rfx_intelligently(self, validated_data: Dict[str, Any], rfx_id: str) -> Dict[str, Any]:
        """
        Ejecuta evaluación inteligente del RFX usando el orquestador.
        
        Args:
            validated_data: Datos del RFX ya validados y limpios
            rfx_id: ID único del RFX para logging
            
        Returns:
            Dict con metadata de evaluación para incluir en RFXProcessed
        """
        # Verificar si la evaluación está habilitada
        if not FeatureFlags.evals_enabled():
            logger.debug(f"🔧 Evaluación inteligente deshabilitada por feature flag para RFX: {rfx_id}")
            return {
                'evaluation_enabled': False,
                'reason': 'Feature flag disabled'
            }
        
        try:
            # Import lazy para evitar circular imports y mejorar startup time
            from backend.services.evaluation_orchestrator import evaluate_rfx_intelligently
            
            logger.info(f"🔍 Iniciando evaluación inteligente para RFX: {rfx_id}")
            start_time = datetime.now()
            
            # Ejecutar evaluación completa
            eval_result = evaluate_rfx_intelligently(validated_data)
            execution_time = (datetime.now() - start_time).total_seconds()
            
            # Log resultados principales
            domain = eval_result['domain_detection']['primary_domain']
            confidence = eval_result['domain_detection']['confidence']
            score = eval_result['consolidated_score']
            quality = eval_result['execution_summary']['overall_quality']
            
            logger.info(f"✅ Evaluación completada para {rfx_id} - Dominio: {domain} ({confidence:.3f}), Score: {score:.3f} ({quality}), Tiempo: {execution_time:.3f}s")
            
            # Extraer recomendaciones críticas y de alta prioridad
            critical_recommendations = [
                {
                    'title': rec['title'],
                    'description': rec['description'],
                    'priority': rec['priority'],
                    'category': rec['category'],
                    'type': rec['type']
                }
                for rec in eval_result['recommendations'] 
                if rec.get('priority') in ['critical', 'high']
            ]
            
            # Log recomendaciones críticas
            if critical_recommendations:
                logger.warning(f"⚠️ RFX {rfx_id} tiene {len(critical_recommendations)} recomendaciones críticas/alta prioridad")
                for rec in critical_recommendations[:2]:  # Log solo las primeras 2
                    logger.warning(f"   - [{rec['priority'].upper()}] {rec['title']}")
            
            # Estructura optimizada de metadata (solo datos esenciales)
            evaluation_metadata = {
                'evaluation_enabled': True,
                'execution_summary': {
                    'consolidated_score': score,
                    'overall_quality': quality,
                    'execution_time_seconds': round(execution_time, 3),
                    'evaluators_executed': eval_result['execution_summary']['evaluators_executed'],
                    'timestamp': datetime.now().isoformat()
                },
                'domain_detection': {
                    'primary_domain': domain,
                    'confidence': confidence,
                    'secondary_domains': eval_result['domain_detection'].get('secondary_domains', [])
                },
                'evaluation_scores': {
                    'generic_score': eval_result['generic_evaluation']['score'],
                    'domain_specific_score': eval_result['domain_specific_evaluation']['score'],
                    'generic_evaluators_count': eval_result['generic_evaluation']['count'],
                    'domain_specific_evaluators_count': eval_result['domain_specific_evaluation']['count']
                },
                'critical_recommendations': critical_recommendations,
                'recommendations_summary': {
                    'total_count': len(eval_result['recommendations']),
                    'critical_count': len([r for r in eval_result['recommendations'] if r.get('priority') == 'critical']),
                    'high_priority_count': len([r for r in eval_result['recommendations'] if r.get('priority') == 'high']),
                    'by_category': {
                        cat: len([r for r in eval_result['recommendations'] if r.get('category') == cat])
                        for cat in set(r.get('category', 'other') for r in eval_result['recommendations'])
                    }
                }
            }
            
            # Log opcional de debugging si está habilitado
            if FeatureFlags.eval_debug_enabled():
                logger.debug(f"🔍 Evaluación detallada para {rfx_id}: {evaluation_metadata}")
            
            return evaluation_metadata
            
        except Exception as e:
            logger.error(f"❌ Error ejecutando evaluación inteligente para {rfx_id}: {str(e)}")
            
            # NO fallar el procesamiento principal - solo log el error
            # Retornar metadata de error para debugging
            return {
                'evaluation_enabled': True,
                'evaluation_error': {
                    'error_message': str(e),
                    'error_type': type(e).__name__,
                    'timestamp': datetime.now().isoformat(),
                    'rfx_id': rfx_id
                },
                'execution_summary': {
                    'consolidated_score': None,
                    'overall_quality': 'evaluation_failed',
                    'execution_time_seconds': 0.0,
                    'evaluators_executed': 0
                }
            }

    def _get_empty_extraction_result(self) -> Dict[str, Any]:
        """Return empty result structure for failed extractions"""
        return {
            "title": "",
            "description": "",
            "email": "",
            "nombre_solicitante": "",
            "productos": [],
            "hora_entrega": "",
            "fecha": "",
            "lugar": "",
            "currency": "USD",
            "texto_original_relevante": "",
            # 🆕 MVP: Requirements fields for compatibility
            "requirements": None,
            "requirements_confidence": 0.0,
            "confidence_score": 0.0  # Legacy compatibility
        }
    
    # ============================================================================
    # 🔍 MÉTODOS DE DEBUGGING Y ESTADÍSTICAS
    # ============================================================================
    
    def get_processing_statistics(self) -> Dict[str, Any]:
        """🆕 Retorna estadísticas de procesamiento para monitoring y debugging"""
        base_stats = self.processing_stats.copy()
        
        # ✅ CAMBIO #3: ELIMINADO - Ya no usamos modular_extractor
        # if hasattr(self, 'modular_extractor'):
        #     extraction_summary = self.modular_extractor.get_extraction_summary()
        #     base_stats.update({
        #         'modular_extractor_stats': extraction_summary.dict(),
        #         'modular_extraction_quality': extraction_summary.extraction_quality,
        #         'modular_chunks_processed': extraction_summary.chunk_count,
        #         'modular_ai_calls': extraction_summary.ai_calls_made,
        #         'modular_retries': extraction_summary.retries_attempted
        #     })
        
        # Calcular ratios y métricas derivadas
        if base_stats['total_documents_processed'] > 0:
            base_stats['fallback_usage_ratio'] = base_stats['fallback_usage_count'] / base_stats['total_documents_processed']
            base_stats['average_chunks_per_document'] = base_stats['chunks_processed'] / base_stats['total_documents_processed']
        else:
            base_stats['fallback_usage_ratio'] = 0.0
            base_stats['average_chunks_per_document'] = 0.0
        
        return base_stats
    
    def reset_processing_statistics(self) -> None:
        """🆕 Resetea estadísticas de procesamiento"""
        self.processing_stats = {
            'total_documents_processed': 0,
            'chunks_processed': 0,
            'average_confidence': 0.0,
            'fallback_usage_count': 0
        }
        
        # ✅ CAMBIO #3: ELIMINADO - Ya no usamos modular_extractor
        # if hasattr(self, 'modular_extractor'):
        #     self.modular_extractor.extraction_stats = {
        #         'chunks_processed': 0,
        #         'ai_calls_made': 0,
        #         'retries_attempted': 0,
        #         'total_processing_time': 0.0
        #     }
        
        logger.info("📊 Processing statistics reset")
    
    def get_debug_mode_status(self) -> Dict[str, Any]:
        """🆕 Retorna estado del modo debug y configuración actual"""
        return {
            'debug_mode_enabled': self.debug_mode,  # ✅ CAMBIO #3: Usar self.debug_mode
            'extraction_strategy': 'intelligent',  # ✅ Siempre inteligente
            'feature_flags': {
                'evals_enabled': FeatureFlags.evals_enabled(),
                'eval_debug_enabled': FeatureFlags.eval_debug_enabled(),
                'meta_prompting_enabled': FeatureFlags.meta_prompting_enabled() if hasattr(FeatureFlags, 'meta_prompting_enabled') else False,
                'vertical_agent_enabled': FeatureFlags.vertical_agent_enabled() if hasattr(FeatureFlags, 'vertical_agent_enabled') else False
            },
            'processing_version': '3.0_simplified',  # ✅ CAMBIO #3: Nueva versión sin extractores
            'available_extractors': [],  # ✅ CAMBIO #3: Ya no hay extractores especializados
            'single_call_extraction': True  # ✅ CAMBIO #3: Todo en una llamada
        }

    # NEW: Multi-file processing
    def _extract_rfx_case_data(
        self,
        rfx_input: RFXInput,
        blobs: List[Dict[str, Any]],
        user_id: str = None,
        organization_id: str = None
    ) -> Dict[str, Any]:
        """
        Pipeline de extracción/resolución SIN persistencia final.
        Retorna datos listos para revisión conversacional y/o guardado posterior.
        """
        logger.info(f"📦 process_rfx_case start: {rfx_input.id} with {len(blobs)} file(s)")
        if user_id:
            logger.info(f"✅ user_id provided for RFX: {user_id}")
        else:
            logger.warning(f"⚠️ No user_id provided for RFX: {rfx_input.id}")
        
        if organization_id:
            logger.info(f"✅ organization_id provided for RFX: {organization_id}")
        else:
            logger.warning(f"⚠️ No organization_id provided for RFX: {rfx_input.id}")
        
        # 🔍 DEBUG: Detailed file logging
        for i, b in enumerate(blobs):
            fname = b.get("filename", f"file_{i}")
            content_size = len(b.get("content", b.get("bytes", b"")))
            logger.info(f"🔍 INPUT FILE {i+1}: '{fname}' ({content_size} bytes)")
        
        # 1) Expand ZIPs
        expanded: List[Dict[str, Any]] = []
        for i, b in enumerate(blobs):
            fname = (b.get("filename") or f"file_{i}").lower()
            content = b.get("content") or b.get("bytes")
            if not content: 
                logger.warning(f"⚠️ EMPTY FILE {i+1}: '{fname}' - skipping")
                continue
            if USE_ZIP and (fname.endswith(".zip") or (content[:2] == b"PK" and not fname.endswith(".docx") and not fname.endswith(".xlsx"))):
                try:
                    with zipfile.ZipFile(io.BytesIO(content)) as zf:
                        for name in zf.namelist():
                            if name.endswith("/"): continue
                            expanded.append({"filename": name, "content": zf.read(name)})
                    logger.info(f"🗜️ ZIP expanded: {fname} → {len(expanded)} internal files")
                except Exception as e:
                    logger.warning(f"⚠️ ZIP expand failed {fname}: {e}")
            else:
                expanded.append({"filename": fname, "content": content})

        text_parts: List[str] = []
        canonical_items: List[Dict[str, Any]] = []
        
        # 2) Per-file routing with DETAILED DEBUG LOGGING
        for file_index, b in enumerate(expanded):
            fname = b["filename"].lower()
            content: bytes = b["content"]
            kind = self._detect_content_type(content, fname)
            logger.info(f"🔎 PROCESSING FILE {file_index+1}: '{fname}' kind={kind} size={len(content)} bytes")
            
            try:
                if kind in ("pdf", "docx", "text"):
                    logger.info(f"📄 Extracting text from {kind.upper()}: {fname}")
                    txt = self._extract_text_from_document(content)
                    logger.info(f"📄 TEXT EXTRACTED from {fname}: {len(txt)} characters")
                    
                    # Fallback OCR if PDF nearly empty
                    if kind == "pdf" and (not txt.strip() or len(re.sub(r"\s+", "", txt)) < 50):
                        logger.info(f"🧠 PDF text too short ({len(txt)} chars), trying OCR for: {fname}")
                        ocr_txt = self._extract_text_with_ocr(content, kind="pdf", filename=fname)
                        if ocr_txt.strip():
                            logger.info(f"🧠 OCR SUCCESS: {fname} → {len(ocr_txt)} characters")
                            txt = ocr_txt
                        else:
                            logger.warning(f"🧠 OCR FAILED for: {fname}")
                    
                    if txt.strip():
                        text_parts.append(f"\n\n### SOURCE: {fname}\n{txt}")
                        logger.info(f"✅ ADDED TO AI CONTEXT: {fname} ({len(txt)} chars)")
                        # Show preview of text to verify content
                        preview = txt[:300].replace('\n', ' ')
                        logger.info(f"📝 CONTENT PREVIEW: {fname} → {preview}...")
                    else:
                        logger.error(f"❌ NO TEXT EXTRACTED from {fname} - file will be ignored!")
                        
                elif kind in ("xlsx", "csv"):
                    logger.info(f"📊 Parsing spreadsheet: {fname}")
                    parsed = self._parse_spreadsheet_items(fname, content)
                    items_count = len(parsed.get('items', []))
                    text_length = len(parsed.get('text', ''))
                    logger.info(f"📊 SPREADSHEET PARSED: {fname} → {items_count} items, {text_length} chars of text")
                    
                    if parsed["items"]: 
                        canonical_items.extend(parsed["items"])
                        logger.info(f"📋 PRODUCTS FOUND in {fname}: {items_count} items")
                        # Log first 3 items for verification
                        for i, item in enumerate(parsed['items'][:3]):
                            logger.info(f"📋 PRODUCT {i+1}: {item}")
                    else:
                        logger.warning(f"⚠️ NO PRODUCTS found in spreadsheet: {fname}")
                    
                    # Always add text for AI metadata extraction
                    if parsed["text"]:  
                        text_parts.append(f"\n\n### SOURCE: {fname}\n{parsed['text']}")
                        preview = parsed['text'][:300].replace('\n', ' ')
                        logger.info(f"📄 SPREADSHEET TEXT ADDED: {fname} → {preview}...")
                    else:
                        # Create summary for AI if we have items but no text
                        if parsed["items"]:
                            summary = f"EXCEL: {len(parsed['items'])} productos encontrados:\n"
                            for item in parsed["items"][:5]:  # First 5 products
                                summary += f"- {item['nombre']}: {item['cantidad']} {item['unidad']}\n"
                            text_parts.append(f"\n\n### SOURCE: {fname}\n{summary}")
                            logger.info(f"📄 SUMMARY CREATED for {fname}: {len(summary)} chars")
                        
                elif kind == "image":
                    logger.info(f"🖼️ Applying OCR to image: {fname}")
                    ocr_txt = self._extract_text_with_ocr(content, kind="image", filename=fname)
                    if ocr_txt.strip():
                        text_parts.append(f"\n\n### SOURCE: {fname} (OCR)\n{ocr_txt}")
                        logger.info(f"✅ IMAGE OCR SUCCESS: {fname} → {len(ocr_txt)} chars")
                        preview = ocr_txt[:300].replace('\n', ' ')
                        logger.info(f"📝 OCR PREVIEW: {fname} → {preview}...")
                    else:
                        logger.error(f"❌ IMAGE OCR FAILED: {fname}")
                else:
                    logger.error(f"❌ UNSUPPORTED FILE TYPE: {fname} (kind={kind})")
                    
            except Exception as e:
                logger.error(f"❌ PROCESSING ERROR for {fname}: {e}")
                import traceback
                logger.error(f"❌ FULL ERROR TRACE: {traceback.format_exc()}")

        combined_text = "\n\n".join(tp for tp in text_parts if tp.strip()) or ""
        
        # 🔍 CRITICAL DEBUG: Show what's being sent to AI
        logger.info(f"📊 FINAL PROCESSING SUMMARY:")
        logger.info(f"📊   - Files processed: {len(expanded)}")
        logger.info(f"📊   - Text parts created: {len(text_parts)}")
        logger.info(f"📊   - Canonical items found: {len(canonical_items)}")
        logger.info(f"📊   - Combined text length: {len(combined_text)} characters")
        
        if combined_text.strip():
            # Show preview of combined text
            preview = combined_text[:500].replace('\n', ' ')
            logger.info(f"📝 COMBINED TEXT PREVIEW (first 500 chars): {preview}...")
            
            # Show text structure
            sources = [part.split('\n')[0] for part in combined_text.split("### SOURCE:") if part.strip()]
            logger.info(f"📄 TEXT SOURCES DETECTED: {sources}")
        else:
            logger.warning(f"⚠️ NO COMBINED TEXT - only canonical items: {len(canonical_items)}")
        
        if not combined_text.strip() and not canonical_items:
            logger.error(f"❌ FATAL: No content extracted from ANY file!")
            raise ValueError("No se pudo extraer texto ni ítems de los archivos proporcionados")

        # 3) AI pipeline
        rfx_input.extracted_content = combined_text
        logger.info(f"🤖 SENDING TO AI: {len(combined_text)} characters of combined text")
        raw_data = self._process_with_ai(combined_text)
        
        # 🛡️ SAFETY CHECK: Handle None raw_data before accessing
        if raw_data is None:
            logger.warning(f"⚠️ raw_data is None, creating minimal structure")
            raw_data = {
                "title": None,
                "description": None,
                "productos": [],
                "email": None,
                "fecha_entrega": None,
                "hora_entrega": None,
                "lugar": None,
                "nombre_solicitante": None,
                "currency": None
            }
        
        # Log AI results
        ai_products_count = len(raw_data.get("productos", []))
        logger.info(f"🤖 AI EXTRACTION RESULTS: {ai_products_count} products found")
        
        if canonical_items:
            logger.info(f"📊 OVERRIDING AI products with canonical spreadsheet products: {len(canonical_items)} items (AI found {ai_products_count})")
            raw_data["productos"] = canonical_items  # Spreadsheet is canonical
        else:
            logger.info(f"📊 USING AI extracted products: {ai_products_count} items")

        # 🛒 Resolver/Enriquecer productos (con o sin catálogo)
        if raw_data.get("productos"):
            # Preparar contexto del RFX para selección inteligente de variantes / bundles.
            rfx_context = {
                'rfx_type': raw_data.get('tipo_evento', 'catering'),
                'description': raw_data.get('descripcion', ''),
                'location': raw_data.get('ubicacion', ''),
                'event_type': raw_data.get('tipo_evento', ''),
                'source_text': (combined_text or "")[:8000],
            }

            if self.catalog_search and organization_id:
                logger.info(f"🛒 Processing {len(raw_data['productos'])} products with catalog pricing...")
                if FeatureFlags.rfx_llm_orchestrator_enabled() and self.rfx_orchestrator_agent:
                    raw_data["productos"] = self._orchestrate_products_with_llm_tools(
                        raw_data["productos"],
                        organization_id,
                        user_id=user_id,
                        rfx_context=rfx_context
                    )
                else:
                    raw_data["productos"] = self._enrich_products_with_catalog(
                        raw_data["productos"],
                        organization_id,
                        user_id=user_id,
                        rfx_context=rfx_context
                    )
            elif self.product_resolution_service:
                logger.info(
                    "🧠 Resolving products with personal/organization fallback context"
                )
                self.product_resolution_service.catalog_search = self.catalog_search
                self.product_resolution_service.rfx_orchestrator_agent = None
                raw_data["productos"] = self.product_resolution_service.resolve_for_chat_products(
                    products=raw_data["productos"],
                    organization_id=organization_id,
                    user_id=user_id,
                    rfx_context=rfx_context,
                )

            # 🔍 DEBUG: Verificar que los productos resueltos tienen precios
            logger.info("🔍 AFTER RESOLUTION - Checking first product:")
            if raw_data["productos"]:
                first_product = raw_data["productos"][0]
                logger.info(f"   📦 Product keys: {list(first_product.keys())}")
                logger.info(f"   💰 costo_unitario: {first_product.get('costo_unitario')}")
                logger.info(f"   💰 precio_unitario: {first_product.get('precio_unitario')}")
                logger.info(f"   💰 unit_cost: {first_product.get('unit_cost')}")
                logger.info(f"   💰 unit_price: {first_product.get('unit_price')}")

        validated_data = self._validate_and_clean_data(raw_data, rfx_input.id)
        
        evaluation_metadata = self._evaluate_rfx_intelligently(validated_data, rfx_input.id)
        rfx_processed = self._create_rfx_processed(validated_data, rfx_input, evaluation_metadata)

        return {
            "rfx_processed": rfx_processed,
            "validated_data": validated_data,
            "evaluation_metadata": evaluation_metadata,
            "combined_text": combined_text,
        }

    def process_rfx_case(self, rfx_input: RFXInput, blobs: List[Dict[str, Any]], user_id: str = None, organization_id: str = None) -> RFXProcessed:
        """Multi-file processing pipeline con persistencia final inmediata (legacy/current flow)."""
        extracted = self._extract_rfx_case_data(
            rfx_input=rfx_input,
            blobs=blobs,
            user_id=user_id,
            organization_id=organization_id,
        )
        rfx_processed: RFXProcessed = extracted["rfx_processed"]
        self._save_rfx_to_database(rfx_processed, user_id=user_id, organization_id=organization_id)
        logger.info(f"✅ process_rfx_case done: {rfx_input.id}")
        return rfx_processed

    def process_rfx_case_preview(
        self,
        rfx_input: RFXInput,
        blobs: List[Dict[str, Any]],
        user_id: str = None,
        organization_id: str = None
    ) -> Dict[str, Any]:
        """
        Procesa RFX para revisión conversacional PREVIA a persistencia.
        No crea registros en rfx_v2 ni rfx_products.
        """
        extracted = self._extract_rfx_case_data(
            rfx_input=rfx_input,
            blobs=blobs,
            user_id=user_id,
            organization_id=organization_id,
        )
        rfx_processed: RFXProcessed = extracted["rfx_processed"]
        validated_data: Dict[str, Any] = extracted["validated_data"] or {}
        combined_text: str = extracted.get("combined_text") or ""

        # Mantener estructura rica de productos para review/chat (incluye specifications/bundles)
        preview_products = validated_data.get("productos") or []

        preview_data = {
            "id": str(rfx_processed.id),
            "title": rfx_processed.title,
            "description": rfx_processed.description,
            "email": rfx_processed.email,
            "requester_name": rfx_processed.requester_name,
            "company_name": rfx_processed.company_name,
            "currency": (rfx_processed.metadata_json or {}).get("validated_currency", "USD"),
            "delivery_date": rfx_processed.delivery_date.isoformat() if rfx_processed.delivery_date else None,
            "delivery_time": rfx_processed.delivery_time.isoformat() if rfx_processed.delivery_time else None,
            "location": rfx_processed.location,
            "requirements": rfx_processed.requirements,
            "requirements_confidence": rfx_processed.requirements_confidence,
            "products": preview_products,
            "metadata_json": rfx_processed.metadata_json or {},
            "source_text": combined_text,
            "rfx_type": (rfx_processed.rfx_type.value if hasattr(rfx_processed.rfx_type, "value") else str(rfx_processed.rfx_type)),
        }

        return {
            "preview_data": preview_data,
            "validated_data": validated_data,
            "evaluation_metadata": extracted.get("evaluation_metadata") or {},
        }

    def _orchestrate_products_with_llm_tools(
        self,
        products: List[Dict[str, Any]],
        organization_id: Optional[str],
        user_id: Optional[str] = None,
        rfx_context: Optional[Dict[str, Any]] = None,
    ) -> List[Dict[str, Any]]:
        """
        Orquesta matching + unidad + pricing con LLM function-calling y tools.
        Fallback automático al flujo legacy si hay error.
        """
        if not products:
            return products

        if not self.product_resolution_service:
            logger.warning("⚠️ ProductResolutionService unavailable, using legacy catalog enrichment")
            return self._enrich_products_with_catalog(products, organization_id, user_id=user_id, rfx_context=rfx_context)

        # Mantener resolver sincronizado con dependencias actuales
        self.product_resolution_service.catalog_search = self.catalog_search
        self.product_resolution_service.rfx_orchestrator_agent = self.rfx_orchestrator_agent

        return self.product_resolution_service.resolve_for_rfx_extraction(
            products=products,
            organization_id=organization_id,
            user_id=user_id,
            rfx_context=rfx_context or {},
            fallback_resolver=lambda fallback_products: self._enrich_products_with_catalog(
                fallback_products,
                organization_id,
                user_id=user_id,
                rfx_context=rfx_context,
            ),
        )

    def _apply_hybrid_bundle_inference(
        self,
        products: List[Dict[str, Any]],
        organization_id: str,
        rfx_context: Dict[str, Any],
    ) -> List[Dict[str, Any]]:
        """
        Modo híbrido:
        1) Si existe producto padre bundle en inventario, usarlo.
        2) Si no existe, crear bundle inferido temporal con pricing seguro.
        """
        if not products:
            return products

        source_text = str(rfx_context.get("source_text") or "")
        lowered = source_text.lower()

        menu_keywords = [
            "menu ejecutivo", "menú ejecutivo",
            "menu saludable", "menú saludable",
            "menu de la semana", "menú de la semana",
            "plan semanal", "semana",
        ]
        weekday_tokens = ["lunes", "martes", "miercoles", "miércoles", "jueves", "viernes"]
        weekday_hits = sum(1 for d in weekday_tokens if d in lowered)
        explicit_menu_signal = any(k in lowered for k in menu_keywords)
        candidate_signal = explicit_menu_signal and weekday_hits >= 2 and len(products) >= 3

        if not candidate_signal:
            return products

        logger.info("🧠 Hybrid mode: weekly/menu signal detected - trying parent bundle resolution")

        parent_variant = None
        if self.catalog_search:
            parent_queries = [
                "Menú Ejecutivo",
                "Menu Ejecutivo",
                "Menú Saludable de la Semana",
                "Menu Saludable de la Semana",
                "Plan Semanal",
            ]
            for query in parent_queries:
                try:
                    variants = self.catalog_search.search_product_variants(
                        query=query,
                        organization_id=organization_id,
                        max_variants=3,
                    )
                    if variants:
                        top = variants[0]
                        if top.get("product_type") in {"complex_bundle", "service_bundle"}:
                            parent_variant = top
                            break
                except Exception as e:
                    logger.warning(f"⚠️ Parent bundle search failed for '{query}': {e}")

        breakdown = self._build_weekly_breakdown(products, source_text)

        if parent_variant:
            price = float(parent_variant.get("unit_price") or 0.0)
            cost = float(parent_variant.get("unit_cost") or 0.0)
            requires_clarification = price <= 0.0
            logger.info(
                f"✅ Hybrid mode using catalog parent bundle: {parent_variant.get('product_name')} "
                f"(price={price}, requires_clarification={requires_clarification})"
            )
            return [
                {
                    "nombre": parent_variant.get("product_name") or "Menú semanal",
                    "descripcion": "Producto bundle detectado automáticamente desde solicitud del cliente",
                    "cantidad": 1.0,
                    "unidad": parent_variant.get("unit") or "plan",
                    "precio_unitario": price,
                    "costo_unitario": cost,
                    "estimated_line_total": round(price, 2),
                    "catalog_match": True,
                    "catalog_product_name": parent_variant.get("product_name"),
                    "catalog_confidence": float(parent_variant.get("confidence") or 0.0),
                    "pricing_source": "catalog_parent_bundle",
                    "requires_clarification": requires_clarification,
                    "bundle_breakdown": breakdown,
                    "specifications": {
                        "is_bundle": True,
                        "inferred_bundle": False,
                        "requires_clarification": requires_clarification,
                        "bundle_breakdown": breakdown,
                    },
                    "notas": "Bundle resuelto con producto padre del inventario",
                }
            ]

        # Fallback híbrido: bundle inferido (sin inventar precio)
        sum_known_sales = 0.0
        sum_known_cost = 0.0
        known_price_items = 0
        for p in products:
            unit_price = float(p.get("precio_unitario") or 0.0)
            unit_cost = float(p.get("costo_unitario") or 0.0)
            qty = float(p.get("cantidad") or 0.0)
            if unit_price > 0 and qty > 0:
                sum_known_sales += unit_price * qty
                sum_known_cost += unit_cost * qty
                known_price_items += 1

        inferred_price = round(sum_known_sales, 2) if known_price_items > 0 else 0.0
        inferred_cost = round(sum_known_cost, 2) if known_price_items > 0 else 0.0
        requires_clarification = known_price_items == 0

        logger.warning(
            "⚠️ Hybrid mode using inferred bundle (not in catalog). "
            f"known_price_items={known_price_items}, inferred_price={inferred_price}"
        )

        return [
            {
                "nombre": "Menú semanal (inferido)",
                "descripcion": "Bundle inferido por IA desde solicitud de menú semanal",
                "cantidad": 1.0,
                "unidad": "plan",
                "precio_unitario": inferred_price,
                "costo_unitario": inferred_cost,
                "estimated_line_total": inferred_price,
                "catalog_match": False,
                "catalog_product_name": None,
                "catalog_confidence": 0.0,
                "pricing_source": "inferred_bundle",
                "requires_clarification": requires_clarification,
                "bundle_breakdown": breakdown,
                "specifications": {
                    "is_bundle": True,
                    "inferred_bundle": True,
                    "requires_clarification": requires_clarification,
                    "pricing_policy": "sum_known_components_or_clarify",
                    "bundle_breakdown": breakdown,
                },
                "notas": (
                    "Bundle inferido automáticamente. "
                    "Requiere confirmación del cliente antes de propuesta final."
                    if requires_clarification
                    else "Bundle inferido automáticamente con suma de componentes conocidos."
                ),
            }
        ]

    def _build_weekly_breakdown(self, products: List[Dict[str, Any]], source_text: str) -> List[Dict[str, Any]]:
        """
        Construye desglose por día (lunes-viernes) usando texto original y fallback por orden.
        """
        day_order = ["lunes", "martes", "miércoles", "jueves", "viernes"]
        text = source_text or ""
        breakdown: List[Dict[str, Any]] = []

        # Intento 1: parseo por patrón "Día: plato"
        for day in day_order:
            pattern = rf"(?is){day}\s*:\s*(?:\n\s*[\*\-]\s*)?([^\n\*]+)"
            m = re.search(pattern, text)
            if m:
                breakdown.append(
                    {
                        "slot": day,
                        "selected": {"name": m.group(1).strip(), "source": "parsed_text"},
                    }
                )

        # Intento 2: fallback por orden de productos extraídos
        if not breakdown:
            for idx, product in enumerate(products[:5]):
                slot = day_order[idx] if idx < len(day_order) else f"dia_{idx+1}"
                breakdown.append(
                    {
                        "slot": slot,
                        "selected": {
                            "name": product.get("nombre") or product.get("product_name") or "producto",
                            "source": "product_order",
                        },
                    }
                )

        return breakdown

    def _detect_content_type(self, content: bytes, filename: str) -> str:
        """Detect file content type from bytes and filename with improved detection"""
        fn = (filename or "").lower()
        
        # Robust PDF detection
        if content.startswith(b"%PDF") or fn.endswith(".pdf"):
            return "pdf"
        
        # DOCX detection (ZIP with specific structure)
        if content[:2] == b"PK":
            if fn.endswith(".docx"):
                return "docx"
            elif fn.endswith(".xlsx"):
                return "xlsx" 
            elif not fn.endswith((".docx", ".xlsx")) and not fn.endswith(".zip"):
                # This could be a DOCX without proper extension
                try:
                    import zipfile
                    with zipfile.ZipFile(io.BytesIO(content)) as zf:
                        # Check for DOCX structure
                        if any("word/" in name for name in zf.namelist()):
                            logger.info(f"🔍 Detected DOCX structure in file without .docx extension: {filename}")
                            return "docx"
                        # Check for XLSX structure
                        elif any("xl/" in name for name in zf.namelist()):
                            logger.info(f"🔍 Detected XLSX structure in file without .xlsx extension: {filename}")
                            return "xlsx"
                except:
                    pass
        
        # Excel and CSV files
        if fn.endswith(".xlsx") or fn.endswith(".xls"):
            return "xlsx"
        if fn.endswith(".csv"):
            return "csv"
        
        # Image files
        image_extensions = [".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".gif", ".webp"]
        if any(fn.endswith(ext) for ext in image_extensions):
            return "image"
        
        # Image magic bytes detection
        if (content.startswith(b"\x89PNG") or 
            content.startswith(b"\xFF\xD8") or  # JPEG
            content.startswith(b"GIF87a") or content.startswith(b"GIF89a") or
            content.startswith(b"RIFF") and b"WEBP" in content[:12]):
            return "image"
        
        # Text files
        if any(fn.endswith(ext) for ext in [".txt", ".md", ".text", ".rtf"]):
            return "text"
        
        # Legacy Word documents
        if fn.endswith(".doc") or content.startswith(b"\xD0\xCF\x11\xE0"):
            return "doc"
        
        # ZIP files
        if fn.endswith(".zip"):
            return "zip"
        
        # Try to detect by content if text-like
        try:
            content.decode('utf-8')
            # If we can decode as UTF-8, probably text
            if len(content) < 50000:  # Small files likely to be text
                return "text"
        except UnicodeDecodeError:
            pass
        
        # Fallback to MIME type guessing
        guessed, _ = mimetypes.guess_type(fn)
        if guessed:
            if "image" in guessed:
                return "image"
            elif "text" in guessed:
                return "text"
            elif "pdf" in guessed:
                return "pdf"
        
        # Final fallback
        logger.warning(f"⚠️ Could not detect content type for {filename}, treating as text")
        return "text"

    def _extract_text_with_ocr(self, file_bytes: bytes, kind: str = "image", filename: Optional[str] = None) -> str:
        """Optional OCR via pytesseract; for PDF, uses pdf2image if available."""
        if not USE_OCR:
            return ""
        try:
            import pytesseract  # type: ignore
            from PIL import Image  # type: ignore
        except Exception as e:
            logger.warning(f"⚠️ OCR unavailable (install pytesseract & pillow): {e}")
            return ""

        if kind == "image":
            try:
                img = Image.open(io.BytesIO(file_bytes))
                return pytesseract.image_to_string(img) or ""
            except Exception as e:
                logger.warning(f"⚠️ OCR image failed: {filename} → {e}")
                return ""

        if kind == "pdf":
            try:
                try:
                    from pdf2image import convert_from_bytes  # type: ignore
                except Exception as e:
                    logger.warning(f"⚠️ pdf2image unavailable for PDF OCR: {e}")
                    return ""
                pages = convert_from_bytes(file_bytes, dpi=200, fmt="png")
                out = []
                for img in pages[:10]:  # safety limit
                    out.append(pytesseract.image_to_string(img) or "")
                return "\n".join(out)
            except Exception as e:
                logger.warning(f"⚠️ OCR PDF failed: {filename} → {e}")
                return ""
        return ""

    def _normalize_date_format(self, date_str: str) -> str:
        """Normalize date formats to YYYY-MM-DD for Pydantic validation"""
        if not date_str:
            return datetime.now().strftime('%Y-%m-%d')
        
        date_str = date_str.strip()
        logger.info(f"🔍 PARSING DATE: '{date_str}'")
        
        # Handle ISO 8601 formats (e.g., "2025-09-06T00:00:00Z", "2025-09-06T00:00:00")
        iso_match = re.match(r'^(\d{4}-\d{1,2}-\d{1,2})T.*$', date_str)
        if iso_match:
            iso_date = iso_match.group(1)
            logger.info(f"📅 EXTRACTED ISO DATE: '{iso_date}' from '{date_str}'")
            try:
                parts = iso_date.split('-')
                year, month, day = int(parts[0]), int(parts[1]), int(parts[2])
                if 1 <= day <= 31 and 1 <= month <= 12 and year > 1900:
                    normalized = f"{year:04d}-{month:02d}-{day:02d}"
                    logger.info(f"✅ ISO DATE NORMALIZED: '{date_str}' → '{normalized}'")
                    return normalized
            except Exception as e:
                logger.warning(f"⚠️ Failed to parse ISO date '{iso_date}': {e}")
        
        # Handle natural language dates in Spanish
        date_natural = self._parse_natural_language_date(date_str)
        if date_natural:
            logger.info(f"✅ NATURAL LANGUAGE DATE: '{date_str}' → '{date_natural}'")
            return date_natural
        
        # Handle DD/MM/YYYY format
        if re.match(r'^\d{1,2}/\d{1,2}/\d{4}$', date_str):
            try:
                parts = date_str.split('/')
                day, month, year = int(parts[0]), int(parts[1]), int(parts[2])
                if 1 <= day <= 31 and 1 <= month <= 12 and year > 1900:
                    normalized = f"{year:04d}-{month:02d}-{day:02d}"
                    logger.info(f"✅ DD/MM/YYYY NORMALIZED: '{date_str}' → '{normalized}'")
                    return normalized
            except Exception as e:
                logger.warning(f"⚠️ Failed to parse DD/MM/YYYY date '{date_str}': {e}")
        
        # Handle DD-MM-YYYY format
        if re.match(r'^\d{1,2}-\d{1,2}-\d{4}$', date_str):
            try:
                parts = date_str.split('-')
                day, month, year = int(parts[0]), int(parts[1]), int(parts[2])
                if 1 <= day <= 31 and 1 <= month <= 12 and year > 1900:
                    normalized = f"{year:04d}-{month:02d}-{day:02d}"
                    logger.info(f"✅ DD-MM-YYYY NORMALIZED: '{date_str}' → '{normalized}'")
                    return normalized
            except Exception as e:
                logger.warning(f"⚠️ Failed to parse DD-MM-YYYY date '{date_str}': {e}")
        
        # Handle YYYY/MM/DD format
        if re.match(r'^\d{4}/\d{1,2}/\d{1,2}$', date_str):
            try:
                parts = date_str.split('/')
                year, month, day = int(parts[0]), int(parts[1]), int(parts[2])
                if 1 <= day <= 31 and 1 <= month <= 12 and year > 1900:
                    normalized = f"{year:04d}-{month:02d}-{day:02d}"
                    logger.info(f"✅ YYYY/MM/DD NORMALIZED: '{date_str}' → '{normalized}'")
                    return normalized
            except Exception as e:
                logger.warning(f"⚠️ Failed to parse YYYY/MM/DD date '{date_str}': {e}")
        
        # If already in YYYY-MM-DD format, validate and return
        if re.match(r'^\d{4}-\d{1,2}-\d{1,2}$', date_str):
            try:
                parts = date_str.split('-')
                year, month, day = int(parts[0]), int(parts[1]), int(parts[2])
                if 1 <= day <= 31 and 1 <= month <= 12 and year > 1900:
                    normalized = f"{year:04d}-{month:02d}-{day:02d}"
                    logger.info(f"✅ YYYY-MM-DD VALIDATED: '{date_str}' → '{normalized}'")
                    return normalized
            except Exception as e:
                logger.warning(f"⚠️ Failed to validate YYYY-MM-DD date '{date_str}': {e}")
        
        # Handle relative dates like "mañana", "today", etc.
        date_str_lower = date_str.lower()
        if date_str_lower in ["hoy", "today"]:
            normalized = datetime.now().strftime('%Y-%m-%d')
            logger.info(f"✅ RELATIVE DATE: '{date_str}' → '{normalized}' (today)")
            return normalized
        elif date_str_lower in ["mañana", "tomorrow"]:
            from datetime import timedelta
            normalized = (datetime.now() + timedelta(days=1)).strftime('%Y-%m-%d')
            logger.info(f"✅ RELATIVE DATE: '{date_str}' → '{normalized}' (tomorrow)")
            return normalized
        
        # Fallback to current date
        current_date = datetime.now().strftime('%Y-%m-%d')
        logger.error(f"❌ COULD NOT PARSE DATE: '{date_str}' - Using current date: {current_date}")
        return current_date

    def _parse_natural_language_date(self, date_str: str) -> str:
        """Parse natural language dates in Spanish like 'octubre 6', '6 de octubre', etc."""
        date_str_lower = date_str.lower().strip()
        
        # Spanish month mapping
        spanish_months = {
            'enero': 1, 'ene': 1,
            'febrero': 2, 'feb': 2,
            'marzo': 3, 'mar': 3,
            'abril': 4, 'abr': 4,
            'mayo': 5, 'may': 5,
            'junio': 6, 'jun': 6,
            'julio': 7, 'jul': 7,
            'agosto': 8, 'ago': 8,
            'septiembre': 9, 'sep': 9, 'sept': 9,
            'octubre': 10, 'oct': 10,
            'noviembre': 11, 'nov': 11,
            'diciembre': 12, 'dic': 12
        }
        
        # English month mapping
        english_months = {
            'january': 1, 'jan': 1,
            'february': 2, 'feb': 2,
            'march': 3, 'mar': 3,
            'april': 4, 'apr': 4,
            'may': 5,
            'june': 6, 'jun': 6,
            'july': 7, 'jul': 7,
            'august': 8, 'aug': 8,
            'september': 9, 'sep': 9, 'sept': 9,
            'october': 10, 'oct': 10,
            'november': 11, 'nov': 11,
            'december': 12, 'dec': 12
        }
        
        # Combine both mappings
        month_mapping = {**spanish_months, **english_months}
        
        current_year = datetime.now().year
        
        # Pattern: "octubre 6", "6 octubre", "octubre 6, 2025", "6 de octubre"
        for month_name, month_num in month_mapping.items():
            if month_name in date_str_lower:
                # Look for day numbers near the month
                words = date_str_lower.replace(',', ' ').replace('de', ' ').split()
                day = None
                year = current_year
                
                # Find day number
                for word in words:
                    if word.isdigit():
                        num = int(word)
                        if 1 <= num <= 31:  # Could be a day
                            day = num
                        elif num > 1900:  # Could be a year
                            year = num
                
                if day:
                    try:
                        # Validate the date
                        import calendar
                        if day <= calendar.monthrange(year, month_num)[1]:
                            result = f"{year:04d}-{month_num:02d}-{day:02d}"
                            logger.info(f"🗓️ PARSED NATURAL DATE: '{date_str}' → month='{month_name}'({month_num}), day={day}, year={year} → '{result}'")
                            return result
                        else:
                            logger.warning(f"⚠️ Invalid day {day} for month {month_num} year {year}")
                    except Exception as e:
                        logger.warning(f"⚠️ Error validating natural date: {e}")
                break
        
        return None

    def _normalize_time_format(self, time_str: str) -> str:
        """Normalize time formats for Pydantic validation"""
        if not time_str:
            return "12:00"
        
        time_str = time_str.strip().lower()
        
        # Handle AM/PM formats
        if 'a.m' in time_str or 'am' in time_str:
            # Extract hour and minute
            time_part = time_str.replace('a.m', '').replace('am', '').strip()
            if ':' in time_part:
                try:
                    hour, minute = time_part.split(':')
                    hour = int(hour)
                    minute = int(minute)
                    # Convert to 24h format
                    if hour == 12:
                        hour = 0  # 12 AM = 00:00
                    return f"{hour:02d}:{minute:02d}"
                except:
                    pass
        elif 'p.m' in time_str or 'pm' in time_str:
            # Extract hour and minute
            time_part = time_str.replace('p.m', '').replace('pm', '').strip()
            if ':' in time_part:
                try:
                    hour, minute = time_part.split(':')
                    hour = int(hour)
                    minute = int(minute)
                    # Convert to 24h format
                    if hour != 12:
                        hour += 12
                    return f"{hour:02d}:{minute:02d}"
                except:
                    pass
        
        # If already in HH:MM format, validate and return
        if ':' in time_str:
            try:
                parts = time_str.split(':')
                if len(parts) == 2:
                    hour, minute = int(parts[0]), int(parts[1])
                    if 0 <= hour <= 23 and 0 <= minute <= 59:
                        return f"{hour:02d}:{minute:02d}"
            except:
                pass
        
        # Fallback
        return "12:00"

    def _enrich_products_with_catalog(
        self, 
        products: List[Dict[str, Any]], 
        organization_id: Optional[str],
        user_id: Optional[str] = None,
        rfx_context: Dict[str, Any] = None
    ) -> List[Dict[str, Any]]:
        """
        Enriquece productos con precios del catálogo usando búsqueda híbrida (SYNC)
        
        Estrategia MEJORADA con selección inteligente de variantes:
        1. Buscar variantes del producto en catálogo (exact → fuzzy → semantic)
        2. Si hay múltiples variantes, usar AI para seleccionar la más apropiada
        3. Si match con confidence >= 0.75, usar precios del catálogo
        4. Si no match, mantener producto original (AI prediction)
        5. Agregar metadata de matching para trazabilidad
        
        Args:
            products: Lista de productos extraídos
            organization_id: ID de la organización
            rfx_context: Contexto del RFX para selección inteligente
        
        Returns:
            Lista de productos enriquecidos con precios de catálogo
        """
        
        if not self.catalog_search:
            logger.warning("⚠️ Catalog search not available, skipping enrichment")
            return products
        
        enriched_products = []
        matches = 0
        misses = 0
        ai_selections = 0
        
        for product in products:
            product_name = product.get('nombre', '')
            
            if not product_name:
                enriched_products.append(product)
                continue
            
            try:
                # 1. Buscar variantes en catálogo (retorna múltiples matches)
                variants = self.catalog_search.search_product_variants(
                    query=product_name,
                    organization_id=organization_id,
                    user_id=user_id,
                    max_variants=5
                )
                
                # 2. Seleccionar la mejor variante
                if variants:
                    if len(variants) > 1:
                        # Regla temporal: elegir siempre la variante con mayor confidence.
                        logger.info(
                            f"🎯 Found {len(variants)} variants for '{product_name}', selecting top confidence"
                        )
                        catalog_match = variants[0]
                        ai_selections += 1
                    else:
                        # Solo 1 variante
                        catalog_match = variants[0]
                else:
                    # No se encontraron variantes
                    catalog_match = None
                
                if catalog_match and catalog_match.get('confidence', 0) >= 0.75:
                    # MATCH ENCONTRADO - usar precios del catálogo
                    matches += 1
                    
                    # Preservar cantidad y unidad del RFX original
                    enriched_product = {
                        **product,
                        'costo_unitario': catalog_match.get('unit_cost'),
                        'precio_unitario': catalog_match.get('unit_price'),
                        
                        # Metadata de matching
                        'catalog_match': True,
                        'catalog_product_name': catalog_match.get('product_name'),
                        'catalog_match_type': catalog_match.get('match_type'),
                        'catalog_confidence': catalog_match.get('confidence'),
                        'pricing_source': 'catalog',
                        
                        # Metadata de selección AI (si aplica)
                        'selection_method': catalog_match.get('selection_method', 'single_match'),
                        'ai_reasoning': catalog_match.get('ai_reasoning'),
                        'variants_count': len(variants) if variants else 0
                    }
                    
                    # Log detallado según método de selección
                    selection_method = catalog_match.get('selection_method', 'single_match')
                    if selection_method == 'ai_intelligent':
                        logger.info(
                            f"✅ AI-selected match: '{product_name}' → "
                            f"'{catalog_match['product_name']}' "
                            f"({len(variants)} variants, confidence={catalog_match['confidence']:.2f}) "
                            f"[cost=${catalog_match.get('unit_cost')}, price=${catalog_match.get('unit_price')}] "
                            f"Reason: {catalog_match.get('ai_reasoning', 'N/A')}"
                        )
                    elif selection_method == 'average_pricing':
                        logger.info(
                            f"✅ Average pricing: '{product_name}' → "
                            f"avg of {len(variants)} variants "
                            f"[cost=${catalog_match.get('unit_cost')}, price=${catalog_match.get('unit_price')}]"
                        )
                    else:
                        logger.info(
                            f"✅ Catalog match: '{product_name}' → "
                            f"'{catalog_match['product_name']}' "
                            f"({catalog_match['match_type']}, confidence={catalog_match['confidence']:.2f}) "
                            f"[cost=${catalog_match.get('unit_cost')}, price=${catalog_match.get('unit_price')}]"
                        )
                    
                    enriched_products.append(enriched_product)
                else:
                    # NO MATCH - mantener producto original
                    misses += 1
                    
                    enriched_product = {
                        **product,
                        'catalog_match': False,
                        'catalog_confidence': catalog_match.get('confidence', 0) if catalog_match else 0,
                        'pricing_source': 'ai_prediction'
                    }
                    
                    logger.info(
                        f"⚠️ No catalog match for: '{product_name}' "
                        f"(confidence={catalog_match.get('confidence', 0):.2f if catalog_match else 0})"
                    )
                    
                    enriched_products.append(enriched_product)
                    
            except Exception as e:
                # Error en búsqueda - mantener producto original
                logger.error(f"❌ Catalog search error for '{product_name}': {e}")
                misses += 1
                
                enriched_product = {
                    **product,
                    'catalog_match': False,
                    'catalog_error': str(e),
                    'pricing_source': 'ai_prediction'
                }
                
                enriched_products.append(enriched_product)
        
        # Actualizar estadísticas
        self.processing_stats['catalog_matches'] += matches
        self.processing_stats['catalog_misses'] += misses
        
        # Log resumen
        total = len(products)
        match_rate = (matches / total * 100) if total > 0 else 0
        
        logger.info(
            f"🛒 CATALOG ENRICHMENT SUMMARY: "
            f"{matches}/{total} matches ({match_rate:.1f}%), "
            f"{misses} misses"
        )
        
        if ai_selections > 0:
            logger.info(f"🤖 AI intelligent selections: {ai_selections}/{matches} matches")
        
        return enriched_products

    def _parse_spreadsheet_items(self, filename: str, content: bytes) -> Dict[str, Any]:
        """Lee XLSX/CSV y devuelve {'items': [...], 'text': 'csv-like'} con mapeo:
        nombre <- Item, cantidad <- Quantity, unidad <- UOM, notes <- Item # (+Description)."""
        try:
            import pandas as pd  # type: ignore
            from io import BytesIO
        except Exception as e:
            logger.warning(f"⚠️ pandas no disponible, saltando parseo de {filename}: {e}")
            return {"items": [], "text": ""}
        try:
            if filename.endswith(".csv"):
                df = pd.read_csv(BytesIO(content))
            else:
                df = pd.read_excel(BytesIO(content))
            if df is None or df.empty:
                logger.warning(f"⚠️ Hoja vacía en {filename}")
                return {"items": [], "text": ""}

            # Normalizar cabeceras
            def _norm(s: str) -> str:
                return str(s).strip().lower()
            cols = {_norm(c): c for c in df.columns}
            def pick(cands: List[str]) -> Optional[str]:
                for c in cands:
                    k = _norm(c)
                    if k in cols: return cols[k]
                    k2 = k.replace(" ", "")
                    if k2 in cols: return cols[k2]
                return None

            itemnum_col = pick(["item #","item#","item no","item nro","item number","numero de item","número de item","id","code","codigo","código"])
            name_col    = pick(["item","producto","product","nombre"])  # ← nombre <- Item
            desc_col    = pick(["description","descripcion","descripción","detalle","item description"])
            qty_col     = pick(["quantity","qty","cantidad","cant","q"])
            unit_col    = pick(["uom","unidad","unidad_medida","unit","units","unidad de medida"])

            logger.info(f"📑 [{filename}] Cols detectadas -> Item#:{itemnum_col} Item:{name_col} Desc:{desc_col} Qty:{qty_col} UOM:{unit_col}")

            def to_int(val) -> int:
                try:
                    if val is None: return 1
                    if isinstance(val, (int, float)): return max(1, int(round(float(val))))
                    s = str(val).strip()
                    if s == "" or s.lower() in {"nan","none","null"}: return 1
                    s = s.replace(" ", "")
                    if "," in s and "." in s:
                        # 1,000.50 -> 1000.50
                        s = s.replace(",", "")
                    else:
                        # 1.000,50 or 25 -> 1 000.50 or 25.0
                        s = s.replace(",", ".")
                    return max(1, int(round(float(s))))
                except Exception:
                    return 1

            def norm_uom(u: Optional[str]) -> str:
                if not u: return "unidades"
                s = str(u).strip().lower().replace(".", "").replace(",", "")
                aliases = {
                    "und":"unidades","un":"unidades","unidad":"unidades","unid":"unidades",
                    "pcs":"unidades","pc":"unidades","pz":"unidades","pza":"unidades",
                    "ea":"unidades","each":"unidades","units":"unidades","unit":"unidades",
                    "set":"set","sets":"set","kg":"kg","kilogram":"kg","kilogramos":"kg",
                    "g":"g","gr":"g","l":"l","lt":"l","liter":"l","litro":"l","litros":"l",
                    "m":"m","mt":"m","caja":"caja","box":"caja","pack":"pack"
                }
                return aliases.get(s, s or "unidades")

            items: List[Dict[str, Any]] = []
            for _, row in df.iterrows():
                nombre = (str(row.get(name_col, "")).strip() if name_col else "")
                if not nombre or nombre.lower() in {"nan","none","null",""}:
                    continue
                cantidad = to_int(row.get(qty_col)) if qty_col else 1
                unidad = norm_uom(str(row.get(unit_col)) if unit_col else None)
                prod: Dict[str, Any] = {"nombre": nombre, "cantidad": cantidad, "unidad": unidad}
                # notes: Item # + Description (si existe y no hay campo dedicated)
                notes_parts = []
                if itemnum_col:
                    v = str(row.get(itemnum_col, "")).strip()
                    if v and v.lower() not in {"nan","none","null",""}:
                        notes_parts.append(f"item_number:{v}")
                if desc_col:
                    d = str(row.get(desc_col, "")).strip()
                    if d and d.lower() not in {"nan","none","null",""}:
                        notes_parts.append(f"desc:{d}")
                if notes_parts:
                    prod["notes"] = " | ".join(notes_parts)
                items.append(prod)

            lines = ["ITEMS_FROM_SPREADSHEET (CANONICAL):", "nombre,cantidad,unidad"]
            for it in items:
                lines.append(f"{it['nombre']},{it['cantidad']},{it['unidad']}")
            logger.info(f"📊 [{filename}] Ítems parseados: {len(items)}")
            return {"items": items, "text": "\n".join(lines)}
        except Exception as e:
            logger.warning(f"⚠️ Falló parseo de hoja {filename}: {e}")
            return {"items": [], "text": ""}
